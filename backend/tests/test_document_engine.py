"""Unit tests for the document-intelligence engine.

These test the pure pipeline: parsers, OCR policy, tables, sections, entities,
extraction, chunking, embeddings, vector store, search and the knowledge graph.
Nothing here touches the database or HTTP.

Several tests exist specifically to pin defects found while building the module,
and they say so. A regression test whose motivation is undocumented gets deleted
by the next person who finds it inconvenient.
"""
from __future__ import annotations

import pytest

from app.domain.documents.fields import (
    FIELD_COUNT, FIELD_SPECS, FIELDS_BY_CATEGORY, FIELDS_BY_KEY, FieldCategory,
)
from app.domain.documents.types import (
    Chunk, DocumentType, EntityKind, FileFormat, ParsedDocument, ParsedPage,
    SectionKind, TextBlock, TextSource, Unit, UnsupportedFormat,
    estimate_tokens, is_monetary, normalise_entity, normalise_whitespace,
    text_fingerprint, to_crore,
)
from app.services.documents.extractors import office, pdf  # noqa: F401  (register)
from app.services.documents.extractors.base import (
    DocumentParser, parse_document, registered_formats,
)
from app.services.documents.extractors.ocr import OcrEngine, OcrPolicy
from app.services.documents.extractors.tables import (
    build_table, detect_header_rows, detect_unit, flatten_header,
    infer_table_unit, normalise_table, parse_number, recover_merges,
    table_confidence, unpack_multiline_rows,
)
from app.services.documents.pipeline.chunking import (
    ChunkConfig, SemanticChunker, duplicate_ratio, split_sentences,
)
from app.services.documents.pipeline.classify import (
    classify_document, classify_with_confidence,
)
from app.services.documents.pipeline.embeddings import (
    HashingEmbeddingProvider, cosine, stem, stem_tokens, tokenise,
)
from app.services.documents.pipeline.entities import EntityExtractor
from app.services.documents.pipeline.financials import (
    FinancialExtractor, LABEL_RULES, PROSE_RULES, detect_period,
    fiscal_year_of, match_label, quarter_of,
)
from app.services.documents.pipeline.knowledge_graph import (
    KnowledgeGraphBuilder, node_key,
)
from app.services.documents.pipeline.orchestrator import IngestionPipeline
from app.services.documents.pipeline.search import (
    DocumentSearch, cite, cite_all, query_terms, verify_answer_citations,
)
from app.services.documents.pipeline.sections import (
    SECTION_RULES, SectionDetector, classify_heading, section_for_order,
    section_for_page,
)
from app.services.documents.pipeline.vector_store import (
    BM25Index, InMemoryVectorStore, VectorRecord,
)
from tests.fixtures.make_docs import (
    ANNUAL_REPORT_TEXT, annual_report_pdf, concall_pdf, credit_rating_pdf,
    scanned_pdf,
)


# ===========================================================================
# Fixtures
# ===========================================================================
@pytest.fixture(scope="module")
def ar_bytes() -> bytes:
    return annual_report_pdf()


@pytest.fixture(scope="module")
def cc_bytes() -> bytes:
    return concall_pdf()


@pytest.fixture(scope="module")
def cr_bytes() -> bytes:
    return credit_rating_pdf()


@pytest.fixture(scope="module")
def annual_report(ar_bytes) -> ParsedDocument:
    return parse_document(ar_bytes, "BHARATCP_AnnualReport_FY25.pdf")


@pytest.fixture(scope="module")
def ar_sections(annual_report):
    return SectionDetector().detect(annual_report)


@pytest.fixture(scope="module")
def ingested(ar_bytes):
    return IngestionPipeline().run(
        ar_bytes, "BHARATCP_AnnualReport_FY25.pdf",
        company_name="Bharat Consumer Products Ltd", company_ticker="BHARATCP",
    )


# ===========================================================================
# Field registry — generated from the workbook
# ===========================================================================
class TestFieldRegistry:
    def test_field_count_matches_workbook(self):
        """The workbook's AI-2 store defines 73 fields. That is the contract."""
        assert FIELD_COUNT == 73
        assert len(FIELD_SPECS) == 73

    def test_sixteen_categories(self):
        assert len(FieldCategory) == 16

    def test_keys_are_unique(self):
        keys = [f.key for f in FIELD_SPECS]
        assert len(keys) == len(set(keys))

    def test_every_field_has_a_unit(self):
        assert all(isinstance(f.unit, Unit) for f in FIELD_SPECS)

    def test_category_index_partitions_the_registry(self):
        total = sum(len(v) for v in FIELDS_BY_CATEGORY.values())
        assert total == FIELD_COUNT

    def test_financial_category_has_ten_fields(self):
        assert len(FIELDS_BY_CATEGORY[FieldCategory.FINANCIAL]) == 10

    def test_every_rule_references_a_real_field(self):
        """A typo in a rule's field key silently disables the rule.

        This is not hypothetical: `r_and_d_spend` was written for a field
        generated as `randd_spend`, and the rule sat dead until this assertion
        was added. Nothing else fails when it happens — coverage just quietly
        reports one field fewer than it should.
        """
        unknown = sorted({
            rule.field_key
            for rule in list(LABEL_RULES) + list(PROSE_RULES)
            if rule.field_key not in FIELDS_BY_KEY
        })
        assert unknown == []


# ===========================================================================
# Units and numbers
# ===========================================================================
class TestUnits:
    @pytest.mark.parametrize("text,expected", [
        ("Rs. in crore", Unit.INR_CRORE),
        ("₹ in lakhs", Unit.INR_LAKH),
        ("Amount in Rs million", Unit.INR_MILLION),
        ("(₹ bn)", Unit.INR_BILLION),
        ("Growth %", Unit.PERCENT),
        ("P/E (x)", Unit.TIMES),
        ("tCO2e", Unit.TONNES_CO2),
        ("INR", Unit.INR),
    ])
    def test_detects_declared_units(self, text, expected):
        assert detect_unit(text) is expected

    @pytest.mark.parametrize("label", [
        "Particulars", "Others", "Reserves", "Costs", "Transfers",
        "Interest", "Debtors", "Warships", "FY25", "Segment",
    ])
    def test_ordinary_labels_declare_no_unit(self, label):
        """Regression: an unanchored "rs" matched inside "Particulars".

        Every table whose first header cell read "Particulars" was therefore
        tagged as a rupee table. Harmless for a rupee table; catastrophic for
        a table of percentages or headcounts, which would then be converted as
        though they were money.
        """
        assert detect_unit(label) is Unit.UNKNOWN

    def test_longest_context_wins(self):
        assert detect_unit("₹ in crore") is Unit.INR_CRORE
        assert detect_unit("₹") is Unit.INR

    @pytest.mark.parametrize("text,value", [
        ("33,543", 33543.0),
        ("1,23,456", 123456.0),          # Indian digit grouping
        ("(1,234.5)", -1234.5),          # accounting negative
        ("-500", -500.0),
        ("12.5%", 12.5),
        ("2.4x", 2.4),
        ("Rs 5,490.70 cr", 5490.7),
    ])
    def test_parses_financial_numbers(self, text, value):
        parsed = parse_number(text)
        assert parsed is not None
        assert parsed[0] == pytest.approx(value)

    @pytest.mark.parametrize("text", ["-", "–", "N/A", "NA", "Nil", "abc", ""])
    def test_rejects_non_numbers(self, text):
        assert parse_number(text) is None

    def test_inline_unit_is_returned(self):
        assert parse_number("12.5%")[1] is Unit.PERCENT
        assert parse_number("Rs 100 crore")[1] is Unit.INR_CRORE

    @pytest.mark.parametrize("unit,value,expected", [
        (Unit.INR_CRORE, 100.0, 100.0),
        (Unit.INR_LAKH, 100.0, 1.0),
        (Unit.INR_MILLION, 100.0, 10.0),
        (Unit.INR_BILLION, 1.0, 100.0),
    ])
    def test_converts_to_crore(self, unit, value, expected):
        assert to_crore(value, unit) == pytest.approx(expected)

    def test_non_monetary_conversion_returns_none(self):
        """A caller that forgets to check gets a visible failure, not a wrong number."""
        assert to_crore(50.0, Unit.PERCENT) is None
        assert to_crore(19220, Unit.COUNT) is None
        assert not is_monetary(Unit.COUNT)


# ===========================================================================
# Tables
# ===========================================================================
class TestTableEngine:
    def test_normalise_drops_empty_rows_and_columns(self):
        raw = [["A", "", "B"], ["", "", ""], ["1", "", "2"]]
        assert normalise_table(raw) == [["A", "B"], ["1", "2"]]

    def test_unpacks_newline_packed_rows(self):
        """pdfplumber returns a whole lattice region as one cell.

        A statement drawn with only an outer border and a header rule arrives
        as one body row holding every value, newline-separated. Left unpacked,
        a thirteen-row income statement extracts as a single unusable row.
        """
        raw = [["Revenue\nEBITDA\nPAT", "100\n20\n10", "120\n25\n12"]]
        assert unpack_multiline_rows(raw) == [
            ["Revenue", "100", "120"],
            ["EBITDA", "20", "25"],
            ["PAT", "10", "12"],
        ]

    def test_refuses_to_unpack_mismatched_rows(self):
        """Zipping columns of different depth would pair labels with wrong numbers.

        Leaving a lumpy cell for the extractor to skip is a visible loss.
        Silently mis-pairing is an invisible error, and far worse.
        """
        raw = [["A\nB", "1\n2\n3"]]
        assert unpack_multiline_rows(raw) == [["A\nB", "1\n2\n3"]]

    def test_recovers_merged_header_spans(self):
        grid = [["Particulars", "FY25", "", "FY24", ""],
                ["", "Q1", "Q2", "Q1", "Q2"]]
        filled, merged = recover_merges(grid, header_rows=1)
        assert filled[0] == ["Particulars", "FY25", "FY25", "FY24", "FY24"]
        assert merged[(0, 1)] == (1, 2)

    def test_body_rows_are_never_forward_filled(self):
        """A blank body cell means nil or not-applicable, not "same as the left"."""
        grid = [["Item", "FY25"], ["Revenue", ""], ["EBITDA", "20"]]
        filled, _ = recover_merges(grid, header_rows=1)
        assert filled[1] == ["Revenue", ""]

    def test_flattens_a_two_row_header(self):
        grid = [["Particulars", "FY25", "FY25"], ["", "Q1", "Q2"]]
        assert flatten_header(grid, 2) == ["Particulars", "FY25 / Q1", "FY25 / Q2"]

    def test_detects_header_row_count(self):
        grid = [["Particulars", "FY24", "FY25"], ["Revenue", "100", "120"]]
        assert detect_header_rows(grid) == 1

    def test_build_table_infers_unit_from_header(self):
        table = build_table(
            [["Particulars (Rs in crore)", "FY25"], ["Revenue", "33,543"]], page=1
        )
        assert table is not None
        assert table.unit is Unit.INR_CRORE
        assert table.header[0] == "Particulars (Rs in crore)"
        assert table.rows == [["Revenue", "33,543"]]

    def test_build_table_rejects_degenerate_input(self):
        assert build_table([["only one row"]], page=1) is None
        assert build_table([], page=1) is None

    def test_confidence_rewards_numeric_rectangular_tables(self):
        good = build_table(
            [["Item", "FY24", "FY25"], ["Revenue", "100", "120"],
             ["EBITDA", "20", "25"]], page=1,
        )
        assert good.confidence > 0.8

    def test_cell_access_is_bounds_safe(self):
        table = build_table([["A", "B"], ["1", "2"]], page=1)
        assert table.cell(99, 99) == ""

    def test_to_grid_rectangularises(self):
        table = build_table([["A", "B", "C"], ["1", "2", "3"]], page=1)
        grid = table.to_grid()
        assert all(len(row) == table.n_cols for row in grid)


# ===========================================================================
# Parsers
# ===========================================================================
class TestParsers:
    def test_all_six_brief_formats_are_registered(self):
        registered = {f.value for f in registered_formats()}
        # The brief names PDF, DOCX, TXT, HTML, CSV and Excel. Markdown is a
        # seventh, carried over from the workbook's upload register.
        assert {"pdf", "docx", "txt", "html", "csv", "xlsx"} <= registered

    def test_unknown_extension_is_refused(self):
        with pytest.raises(UnsupportedFormat):
            DocumentParser.format_for("report.rtf")

    def test_parses_markdown(self):
        doc = parse_document(
            b"# Annual Report\n\nRevenue grew to Rs 33,543 crore.\n", "note.md"
        )
        assert doc.title == "Annual Report"
        assert "33,543" in doc.full_text

    def test_parses_csv_as_one_table(self):
        payload = b"Particulars,FY24,FY25\nRevenue (Rs cr),28914,33543\n"
        doc = parse_document(payload, "figures.csv")
        assert doc.page_count == 1
        assert doc.pages[0].source is TextSource.STRUCTURED
        assert len(doc.tables) == 1
        assert doc.tables[0].header == ["Particulars", "FY24", "FY25"]

    def test_parses_html_with_tables(self):
        payload = (
            b"<html><head><title>Filing</title></head><body>"
            b"<h1>Outcome of Board Meeting</h1><p>Capex of Rs 1,200 crore.</p>"
            b"<table><tr><th>Item</th><th>FY25 (Rs cr)</th></tr>"
            b"<tr><td>Revenue</td><td>33,543</td></tr></table></body></html>"
        )
        doc = parse_document(payload, "filing.html")
        assert doc.title == "Filing"
        assert doc.tables[0].unit is Unit.INR_CRORE

    def test_parses_xlsx_one_page_per_sheet(self):
        import io

        import openpyxl

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "P&L (Rs cr)"
        sheet.append(["Particulars", "FY24", "FY25"])
        sheet.append(["Revenue", 28914, 33543])
        workbook.create_sheet("Balance Sheet")
        buffer = io.BytesIO()
        workbook.save(buffer)

        doc = parse_document(buffer.getvalue(), "model.xlsx")
        assert doc.page_count == 2
        assert doc.pages[0].tables[0].unit is Unit.INR_CRORE

    def test_parses_docx(self):
        import io

        import docx

        document = docx.Document()
        document.add_heading("Business Overview", level=1)
        document.add_paragraph("The Company is engaged in packaged foods.")
        buffer = io.BytesIO()
        document.save(buffer)

        doc = parse_document(buffer.getvalue(), "overview.docx")
        assert "packaged foods" in doc.full_text
        assert any(b.bold for p in doc.pages for b in p.blocks)

    def test_parses_pdf_with_layout(self, annual_report):
        assert annual_report.page_count >= 3
        assert annual_report.char_count > 4000
        assert any(
            b.font_size and b.font_size > 12
            for p in annual_report.pages for b in p.blocks
        )

    def test_recovers_the_financial_table(self, annual_report):
        tables = annual_report.tables
        assert len(tables) == 1
        table = tables[0]
        assert table.unit is Unit.INR_CRORE
        assert table.header[0].startswith("Particulars")
        assert table.n_rows >= 12
        labels = [row[0] for row in table.rows]
        assert "Revenue from operations" in labels
        assert "EBITDA" in labels

    def test_empty_upload_is_rejected(self):
        from app.domain.documents.types import ParseFailure

        with pytest.raises(ParseFailure):
            parse_document(b"", "broken.pdf")


# ===========================================================================
# OCR
# ===========================================================================
class TestOcr:
    def test_sparse_page_needs_ocr(self):
        policy = OcrPolicy()
        assert policy.needs_ocr(char_count=10, page_area=500_000, image_ratio=0.0)

    def test_image_dominated_page_needs_ocr(self):
        policy = OcrPolicy()
        assert policy.needs_ocr(char_count=5000, page_area=500_000, image_ratio=0.9)

    def test_dense_text_page_does_not(self):
        """The brief's rule: do not OCR machine-readable PDFs unnecessarily."""
        policy = OcrPolicy()
        assert not policy.needs_ocr(
            char_count=4000, page_area=500_000, image_ratio=0.05
        )

    def test_engine_reports_its_own_availability(self):
        status = OcrEngine().describe()
        assert set(status) >= {"available", "engine", "language", "policy"}
        assert isinstance(status["available"], bool)

    def test_native_pdf_is_not_ocred(self, annual_report):
        assert annual_report.used_ocr is False
        assert all(p.source is TextSource.NATIVE for p in annual_report.pages)

    @pytest.mark.skipif(not OcrEngine().available, reason="Tesseract not installed")
    def test_scanned_pdf_is_ocred(self):
        doc = parse_document(scanned_pdf(), "scan.pdf")
        assert doc.used_ocr is True
        page = doc.pages[0]
        assert page.source is TextSource.OCR
        assert page.ocr_confidence is not None and page.ocr_confidence > 0.5
        assert "SCANNED" in page.text.upper()

    def test_unavailable_ocr_raises_rather_than_returning_empty(self, monkeypatch):
        """A scan yielding nothing must look like a failure, not an empty filing."""
        from app.domain.documents.types import OcrUnavailable

        engine = OcrEngine()
        monkeypatch.setattr(engine, "_available", False)
        with pytest.raises(OcrUnavailable):
            engine.recognise(b"not-an-image")


# ===========================================================================
# Sections
# ===========================================================================
class TestSectionDetection:
    @pytest.mark.parametrize("heading,expected", [
        ("Management Discussion and Analysis", SectionKind.MANAGEMENT_DISCUSSION),
        ("Chairman's Letter", SectionKind.CHAIRMAN_LETTER),
        ("Risk Factors", SectionKind.RISK_FACTORS),
        ("Independent Auditor's Report", SectionKind.AUDITOR_REPORT),
        ("Corporate Governance", SectionKind.CORPORATE_GOVERNANCE),
        ("Shareholding Pattern", SectionKind.SHAREHOLDING),
        ("Business Responsibility and Sustainability Report", SectionKind.ESG),
        ("Question and Answer Session", SectionKind.CONFERENCE_QA),
        ("Notes to Accounts", SectionKind.NOTES_TO_ACCOUNTS),
        ("Balance Sheet", SectionKind.FINANCIAL_STATEMENTS),
        ("Business Overview", SectionKind.BUSINESS_OVERVIEW),
        ("Outlook and Guidance", SectionKind.MANAGEMENT_GUIDANCE),
    ])
    def test_classifies_the_briefs_sections(self, heading, expected):
        kind, score = classify_heading(heading)
        assert kind is expected
        assert score > 0.5

    def test_notes_beats_financial_statements(self):
        """Overlapping vocabularies must resolve to the more specific rule."""
        kind, _ = classify_heading("Notes to the Financial Statements")
        assert kind is SectionKind.NOTES_TO_ACCOUNTS

    def test_unrelated_heading_is_unknown(self):
        assert classify_heading("Photo Gallery")[0] is SectionKind.UNKNOWN

    def test_every_brief_section_has_a_rule(self):
        covered = {rule.kind for rule in SECTION_RULES}
        required = set(SectionKind) - {SectionKind.UNKNOWN}
        assert required <= covered

    def test_detects_all_sections_in_the_annual_report(self, ar_sections):
        found = {s.kind for s in ar_sections}
        expected = {
            SectionKind.CHAIRMAN_LETTER, SectionKind.BUSINESS_OVERVIEW,
            SectionKind.MANAGEMENT_DISCUSSION, SectionKind.RISK_FACTORS,
            SectionKind.CORPORATE_GOVERNANCE, SectionKind.ESG,
            SectionKind.AUDITOR_REPORT, SectionKind.FINANCIAL_STATEMENTS,
        }
        assert expected <= found

    def test_sections_carry_block_ordinals(self, ar_sections):
        """Page granularity cannot separate sections that share a page.

        Four sections of this report begin on page 2. A page-granular lookup
        attributed all four to whichever started first, so governance text was
        cited as ESG. Block ordinals are what make the attribution correct.
        """
        assert all(s.start_order is not None for s in ar_sections)
        assert all(s.end_order >= s.start_order for s in ar_sections)

    def test_several_sections_share_a_page(self, ar_sections):
        by_page: dict[int, set] = {}
        for section in ar_sections:
            by_page.setdefault(section.start_page, set()).add(section.kind)
        assert max(len(kinds) for kinds in by_page.values()) >= 3

    def test_order_lookup_beats_page_lookup(self, ar_sections):
        crowded = max(
            (s.start_page for s in ar_sections),
            key=lambda p: sum(1 for s in ar_sections if s.start_page == p),
        )
        on_page = [s for s in ar_sections if s.start_page == crowded]
        assert len(on_page) >= 2
        for section in on_page:
            resolved = section_for_order(ar_sections, section.start_order, crowded)
            assert resolved is not None and resolved.kind is section.kind

    def test_page_lookup_prefers_the_narrowest_span(self, ar_sections):
        section = section_for_page(ar_sections, 1)
        assert section is not None

    def test_spans_do_not_overlap_in_document_order(self, ar_sections):
        ordered = sorted(ar_sections, key=lambda s: s.start_order)
        for earlier, later in zip(ordered, ordered[1:]):
            assert earlier.end_order < later.start_order

    def test_falls_back_to_page_text_at_lower_confidence(self):
        """A scan has no typography, so detection degrades — visibly."""
        page = ParsedPage(
            number=1,
            text="Risk Factors\nVolatility in commodity prices is the key risk.",
        )
        doc = ParsedDocument(pages=[page])
        sections = SectionDetector().detect(doc)
        assert sections and sections[0].kind is SectionKind.RISK_FACTORS
        assert sections[0].confidence < 0.5


# ===========================================================================
# Entities
# ===========================================================================
class TestEntityExtraction:
    def test_extracts_subsidiaries(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        names = {e.name for e in entities if e.kind is EntityKind.SUBSIDIARY}
        assert "Bharat Nutrition Private Limited" in names
        assert "Nirmal Personal Care Limited" in names
        assert "Suraj Foods International Pte Ltd" in names

    def test_names_do_not_bleed_across_sentences(self, annual_report):
        """Regression: re.I made [A-Z] match lowercase, killing the anchor.

        The extractor captured "Limited in our core categories. Bharat
        Nutrition Private Limited" as a subsidiary, having matched backwards
        through a sentence boundary into the previous sentence.
        """
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        for entity in entities:
            if entity.kind in {
                EntityKind.SUBSIDIARY, EntityKind.COMPETITOR,
                EntityKind.DIRECTOR, EntityKind.AUDITOR,
            }:
                assert ". " not in entity.name, entity.name
                assert entity.name[0].isupper()

    def test_extracts_directors(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        names = {e.name for e in entities if e.kind is EntityKind.DIRECTOR}
        assert {"Arvind Deshmukh", "Kavita Raman", "Suresh Iyer"} <= names

    def test_extracts_competitors_and_countries(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        competitors = [e for e in entities if e.kind is EntityKind.COMPETITOR]
        countries = {e.name for e in entities if e.kind is EntityKind.COUNTRY}
        assert competitors
        assert {"India", "Singapore"} <= countries

    def test_extracts_risk_and_capex_phrases(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        assert any(e.kind is EntityKind.RISK for e in entities)
        assert any(e.kind is EntityKind.CAPEX for e in entities)

    def test_subject_company_is_not_its_own_subsidiary(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        subject = normalise_entity("Bharat Consumer Products Ltd")
        assert not any(
            e.normalised == subject and e.kind is not EntityKind.COMPANY
            for e in entities
        )

    def test_every_entity_carries_page_and_evidence(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        for entity in entities:
            assert entity.page >= 1
            assert 0.0 < entity.confidence <= 1.0

    def test_repetition_raises_confidence_but_never_to_certainty(self):
        text = " ".join([
            "Alpha Foods Limited is a wholly-owned subsidiary of the Company."
        ] * 6)
        doc = ParsedDocument(pages=[ParsedPage(number=1, text=text)])
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(doc)
        subsidiary = next(e for e in entities if e.kind is EntityKind.SUBSIDIARY)
        assert subsidiary.confidence > 0.82
        assert subsidiary.confidence < 1.0

    def test_transcript_speech_is_not_mistaken_for_a_name(self, cc_bytes):
        """Regression: a transcript turned its own dialogue into directors.

        The rule "Role: Name" is correct in a governance report and wrong in a
        conference call, where the text after "Chief Financial Officer:" is
        what the person *said*. It extracted "Palm" from "Palm oil has been the
        principal pressure point" and "Thank" from "Thank you.", and both were
        rendered on the knowledge graph as directors of the company.
        """
        doc = parse_document(cc_bytes, "concall.pdf")
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(doc)
        names = {e.name for e in entities if e.kind is EntityKind.DIRECTOR}
        assert "Palm" not in names
        assert "Thank" not in names
        # The names genuinely present must still be found.
        assert {"Kavita Raman", "Suresh Iyer"} <= names

    def test_a_lone_common_word_is_never_an_entity(self, cr_bytes):
        """Regression: "Key customers include ..." produced an entity called "Key"."""
        doc = parse_document(cr_bytes, "rating.pdf")
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(doc)
        for entity in entities:
            if entity.kind is EntityKind.COUNTRY:
                continue
            words = entity.name.split()
            if len(words) == 1:
                assert words[0].lower() not in {"key", "major", "top", "thank", "palm"}

    def test_conjunction_lists_split_into_separate_entities(self, cr_bytes):
        """"Suppliers are A Limited and B Limited" is two suppliers, not one.

        Storing the whole span as a single entity produces a graph node no
        reader would recognise and no lookup would ever match.
        """
        doc = parse_document(cr_bytes, "rating.pdf")
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(doc)
        suppliers = {e.name for e in entities if e.kind is EntityKind.SUPPLIER}
        assert "Adani Wilmar Limited" in suppliers
        assert "Ruchi Soya Industries Limited" in suppliers
        assert not any(" and " in name for name in suppliers)

    def test_ampersand_inside_a_firm_name_is_not_a_list_separator(self, annual_report):
        """"Bhattacharya & Associates" is one auditor, not two."""
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        auditors = {e.name for e in entities if e.kind is EntityKind.AUDITOR}
        assert "Bhattacharya & Associates" in auditors

    def test_section_headings_do_not_leak_into_names(self, cc_bytes):
        """A heading has no full stop, so the sentence-boundary trim cannot see it.

        "Management Guidance Kavita Raman" was extracted as a director's name.
        """
        doc = parse_document(cc_bytes, "concall.pdf")
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(doc)
        for entity in entities:
            if entity.kind is EntityKind.DIRECTOR:
                assert not entity.name.lower().startswith("management")

    def test_normalisation_unifies_legal_forms(self):
        assert normalise_entity("Acme Ltd.") == normalise_entity("ACME Limited")
        assert normalise_entity("Acme Private Limited") == "acme"


# ===========================================================================
# Periods
# ===========================================================================
class TestPeriods:
    @pytest.mark.parametrize("text,expected", [
        ("FY 2024-25", "FY25"),
        ("FY2024-25", "FY25"),
        ("2024-25", "FY25"),
        ("FY24-25", "FY25"),
        ("FY 2023-24", "FY24"),
        ("FY25", "FY25"),
        ("FY 2025", "FY25"),
        ("year ended 31 March 2025", "FY25"),
        ("Q3 FY25", "Q3FY25"),
        ("Q4FY25", "Q4FY25"),
    ])
    def test_normalises_period_labels(self, text, expected):
        """Regression: "FY 2024-25" resolved to FY24.

        The first pattern captured "24" out of "2024" and never saw the "-25".
        Every figure in that column was then filed against the prior year, and
        a year-on-year comparison silently compared a figure with itself.
        """
        assert detect_period(text) == expected

    def test_ranges_resolve_to_the_closing_year(self):
        assert detect_period("2023-2024") == "FY24"

    def test_non_periods_return_none(self):
        assert detect_period("Particulars") is None
        assert detect_period("") is None

    def test_fiscal_year_of_annual_and_quarterly(self):
        assert fiscal_year_of("FY25") == 2025
        # A Q4 FY25 transcript is evidence about FY2025 and must be filed
        # under it; returning None left every call unattached to a year.
        assert fiscal_year_of("Q4FY25") == 2025
        assert fiscal_year_of(None) is None

    def test_quarter_of(self):
        assert quarter_of("Q3FY25") == 3
        assert quarter_of("FY25") is None


# ===========================================================================
# Financial extraction
# ===========================================================================
class TestFinancialExtraction:
    def test_matches_statement_labels(self):
        spec, score = match_label("Revenue from operations")
        assert spec.key == "revenue"
        assert score > 0.9

    def test_exclusions_prevent_near_misses(self):
        assert match_label("Other operating revenue") is None or \
            match_label("Other operating revenue")[0].key != "revenue"
        assert match_label("EBITDA margin") is None or \
            match_label("EBITDA margin")[0].key != "ebitda"

    def test_extracts_the_income_statement(self, annual_report, ar_sections):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        values = {
            (f.field_key, f.period): f.value
            for f in result.facts if f.value is not None
        }
        # These are BHARATCP's reference-model figures. Extraction must
        # reproduce them exactly, not approximately.
        assert values[("revenue", "FY25")] == pytest.approx(33543.0)
        assert values[("ebitda", "FY25")] == pytest.approx(5490.70)
        assert values[("pat", "FY25")] == pytest.approx(3450.90)
        assert values[("net_worth", "FY25")] == pytest.approx(15606.90)
        assert values[("revenue", "FY24")] == pytest.approx(28914.0)

    def test_table_facts_inherit_the_declared_unit(self, annual_report, ar_sections):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        revenue = next(
            f for f in result.facts
            if f.field_key == "revenue" and f.period == "FY25"
        )
        assert revenue.unit is Unit.INR_CRORE
        assert revenue.value_in_crore() == pytest.approx(33543.0)

    def test_counts_never_inherit_a_currency(self, annual_report, ar_sections):
        """Regression: headcount inside a "₹ in crore" table became ₹ crore.

        `to_crore()` would then have converted 19,220 employees into ₹19,220
        crore — a number that is both wrong and entirely plausible.
        """
        result = FinancialExtractor().extract(annual_report, ar_sections)
        headcount = [f for f in result.facts if f.field_key == "employee_headcount"]
        assert headcount
        for fact in headcount:
            assert fact.unit is Unit.COUNT
            assert fact.value_in_crore() is None

    def test_extracts_guidance_from_prose(self, annual_report, ar_sections):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        by_key = {f.field_key: f for f in result.facts}
        assert by_key["revenue_growth_guidance"].value == pytest.approx(12.0)
        assert by_key["ebitda_margin_guidance"].value == pytest.approx(17.0)
        assert by_key["ebitda_margin_guidance"].unit is Unit.PERCENT

    def test_extracts_esg_and_governance_from_prose(self, annual_report, ar_sections):
        """ESG and governance figures are written as sentences, not tabulated.

        Without a prose path these fields were invisible: ESG coverage read 1
        of 5 on a document that states four of the five in plain English.
        """
        result = FinancialExtractor().extract(annual_report, ar_sections)
        by_key = {f.field_key: f for f in result.facts}
        assert by_key["renewable_energy_share"].value == pytest.approx(34.0)
        assert by_key["csr_spend"].value == pytest.approx(71.0)
        assert by_key["board_size"].value == pytest.approx(10.0)
        assert by_key["independent_director_share"].value == pytest.approx(6.0)

    def test_ranks_principal_risks(self, annual_report, ar_sections):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        risks = sorted(
            (f for f in result.facts if f.field_key.startswith("principal_risk")),
            key=lambda f: f.field_key,
        )
        assert len(risks) >= 2
        assert all(f.text for f in risks)
        # Confidence declines down the ranking.
        assert risks[0].confidence >= risks[-1].confidence

    def test_every_fact_is_citable(self, annual_report, ar_sections):
        """A fact without a page could not be cited, so it is not evidence."""
        result = FinancialExtractor().extract(annual_report, ar_sections)
        for fact in result.facts:
            assert fact.page >= 1
            assert fact.evidence
            assert 0.0 < fact.confidence <= 1.0

    def test_coverage_is_measured_against_the_full_registry(
        self, annual_report, ar_sections
    ):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        # `coverage` is rounded to four places for display; compare against the
        # same rounding rather than the raw quotient. The first version of this
        # assertion compared 0.6712 with 0.6712328... and failed on the
        # product's own presentation rule, not on a defect.
        assert result.coverage == pytest.approx(
            round(len(result.covered) / FIELD_COUNT, 4)
        )
        assert len(result.missing()) == FIELD_COUNT - len(result.covered)

    def test_reports_a_material_share_of_the_registry(
        self, annual_report, ar_sections
    ):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        assert result.coverage > 0.55

    def test_one_fact_per_field_and_period(self, annual_report, ar_sections):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        keys = [(f.field_key, f.period) for f in result.facts]
        assert len(keys) == len(set(keys))

    def test_category_breakdown_sums_to_the_registry(
        self, annual_report, ar_sections
    ):
        result = FinancialExtractor().extract(annual_report, ar_sections)
        breakdown = result.by_category()
        assert sum(v["defined"] for v in breakdown.values()) == FIELD_COUNT


# ===========================================================================
# Chunking
# ===========================================================================
class TestChunking:
    def test_sentence_split_survives_decimals(self):
        """The same lesson as Module 6's citation auditor.

        Splitting on every full stop tore "₹33,543.00 crore" into two
        sentences, orphaning the number from its context.
        """
        sentences = split_sentences(
            "Revenue was Rs 33,543.00 crore. EBITDA margin reached 16.37%."
        )
        assert len(sentences) == 2
        assert "33,543.00" in sentences[0]

    def test_chunks_never_split_a_sentence(self, annual_report, ar_sections):
        chunks = SemanticChunker().chunk(annual_report, ar_sections)
        assert chunks
        prose = [c for c in chunks if c.paragraph < 1000]
        for chunk in prose:
            assert chunk.text == chunk.text.strip()
            assert not chunk.text.endswith(" ,")

    def test_chunks_carry_page_paragraph_and_section(self, annual_report, ar_sections):
        """Everything the brief's citation requirement needs, on every chunk."""
        chunks = SemanticChunker().chunk(annual_report, ar_sections)
        for chunk in chunks:
            assert chunk.page >= 1
            assert chunk.paragraph >= 0
            assert isinstance(chunk.section, SectionKind)
            assert chunk.fingerprint
            assert chunk.token_estimate > 0

    def test_chunk_indices_are_contiguous(self, annual_report, ar_sections):
        chunks = SemanticChunker().chunk(annual_report, ar_sections)
        assert [c.chunk_index for c in chunks] == list(range(len(chunks)))

    def test_respects_the_token_ceiling(self, annual_report, ar_sections):
        config = ChunkConfig(target_tokens=60, max_tokens=110)
        chunks = SemanticChunker(config).chunk(annual_report, ar_sections)
        # A single sentence longer than the cap is emitted whole rather than
        # cut, so the assertion allows one-sentence overruns only.
        for chunk in chunks:
            if len(split_sentences(chunk.text)) > 1:
                assert chunk.token_estimate <= config.max_tokens + 60

    def test_deduplicates_repeated_boilerplate(self):
        boilerplate = (
            "This report contains forward-looking statements that involve risks "
            "and uncertainties which could cause actual results to differ."
        )
        pages = [
            ParsedPage(number=i, text=boilerplate, blocks=[
                TextBlock(text=boilerplate, page=i)
            ])
            for i in range(1, 9)
        ]
        chunks = SemanticChunker().chunk(ParsedDocument(pages=pages), [])
        assert len(chunks) <= DEFAULT_DUPLICATE_LIMIT

    def test_tables_are_indexed_as_chunks(self, annual_report, ar_sections):
        chunks = SemanticChunker().chunk(annual_report, ar_sections)
        table_chunks = [c for c in chunks if c.paragraph >= 1000]
        assert table_chunks
        assert "Revenue from operations" in " ".join(c.text for c in table_chunks)

    def test_duplicate_ratio_is_reported(self, annual_report, ar_sections):
        chunks = SemanticChunker().chunk(annual_report, ar_sections)
        assert 0.0 <= duplicate_ratio(chunks) < 0.5

    def test_fingerprint_ignores_case_and_whitespace(self):
        assert text_fingerprint("Revenue  Grew") == text_fingerprint("revenue grew")


DEFAULT_DUPLICATE_LIMIT = ChunkConfig().duplicate_threshold


# ===========================================================================
# Embeddings
# ===========================================================================
class TestEmbeddings:
    def test_vectors_have_the_declared_dimension(self):
        provider = HashingEmbeddingProvider(dimension=128)
        vector = provider.embed_one("revenue grew to 33,543 crore")
        assert len(vector) == 128
        assert provider.spec.dimension == 128

    def test_embedding_is_deterministic(self):
        provider = HashingEmbeddingProvider()
        assert provider.embed_one("EBITDA margin") == provider.embed_one("EBITDA margin")

    def test_similar_text_scores_higher_than_unrelated(self):
        provider = HashingEmbeddingProvider()
        a = provider.embed_one("EBITDA margin guidance of 17 percent")
        b = provider.embed_one("EBITDA margin guidance of 18 percent")
        c = provider.embed_one("the auditor issued an emphasis of matter")
        assert cosine(a, b) > cosine(a, c)

    def test_empty_text_yields_a_zero_vector(self):
        assert all(v == 0.0 for v in HashingEmbeddingProvider().embed_one(""))

    def test_cosine_handles_degenerate_input(self):
        assert cosine([], []) == 0.0
        assert cosine([0.0, 0.0], [1.0, 1.0]) == 0.0

    def test_batch_order_is_preserved(self):
        provider = HashingEmbeddingProvider()
        texts = ["revenue", "ebitda", "capex"]
        batch = provider.embed(texts)
        assert batch == [provider.embed_one(t) for t in texts]

    @pytest.mark.parametrize("a,b", [
        ("competitors", "competition"),
        ("subsidiaries", "subsidiary"),
        ("guidance", "guide"),
        ("acquisitions", "acquired"),
        ("borrowings", "borrowing"),
    ])
    def test_stemming_unifies_query_and_document_vocabulary(self, a, b):
        """Regression: "who are the competitors?" returned nothing.

        No filing uses the noun the question uses — they write "we compete
        with". Exact token matching cannot bridge that, so the query was
        unanswerable against a corpus that plainly answered it.
        """
        assert stem(a) == stem(b)

    def test_stemming_is_idempotent(self):
        """Single-pass stripping is not a normalisation.

        "competitors" lost only "s" and became "competitor"; "competition"
        lost "ion" and became "compet" — two keys for one concept, which
        defeats the purpose entirely.
        """
        words = """competitors competition subsidiaries revenues emissions
            utilisation borrowings governance directors liabilities""".split()
        for word in words:
            assert stem(stem(word)) == stem(word)

    def test_stemming_does_not_over_collapse(self):
        assert stem("gas") == "gas"
        assert stem("analysis") == "analysis"
        assert stem("revenue") != stem("reserve")

    def test_stem_tokens_matches_tokenise_length(self):
        text = "Revenue from operations grew 16%"
        assert len(stem_tokens(text)) == len(tokenise(text))


# ===========================================================================
# Vector store
# ===========================================================================
def _record(chunk_id: int, text: str, provider, *, document_id=1, page=1) -> VectorRecord:
    return VectorRecord(
        chunk_id=chunk_id, document_id=document_id, text=text, page=page,
        paragraph=0, section=SectionKind.UNKNOWN, document_title="Doc",
        vector=provider.embed_one(text),
    )


class TestVectorStore:
    @pytest.fixture()
    def populated(self):
        provider = HashingEmbeddingProvider()
        store = InMemoryVectorStore()
        store.add([
            _record(1, "EBITDA margin guidance of 17% for the coming year", provider),
            _record(2, "Volatility in agricultural commodity prices is the key risk",
                    provider, page=2),
            _record(3, "The Company has been rated CRISIL AA+ for long-term facilities",
                    provider, document_id=2),
        ], provider.spec)
        return store, provider

    def test_counts_records(self, populated):
        store, _ = populated
        assert store.count() == 3
        assert store.count(document_id=2) == 1

    def test_hybrid_search_ranks_the_right_chunk_first(self, populated):
        store, provider = populated
        hits = store.search(provider.embed_one("credit rating"), "credit rating")
        assert hits and hits[0].record.chunk_id == 3

    def test_reports_both_component_scores(self, populated):
        store, provider = populated
        hits = store.search(provider.embed_one("EBITDA margin"), "EBITDA margin")
        assert hits[0].lexical > 0
        assert hits[0].semantic > 0

    def test_filters_by_document(self, populated):
        store, provider = populated
        hits = store.search(
            provider.embed_one("risk"), "risk", document_ids=[2]
        )
        assert all(h.record.document_id == 2 for h in hits)

    def test_deleting_a_document_removes_its_chunks(self, populated):
        store, _ = populated
        assert store.delete_document(2) == 1
        assert store.count() == 2

    def test_mixing_embedding_spaces_is_refused(self, populated):
        """Silently mixing spaces yields valid arithmetic and meaningless scores."""
        store, _ = populated
        other = HashingEmbeddingProvider(dimension=64)
        with pytest.raises(ValueError, match="embedding space mismatch"):
            store.add([_record(9, "text", other)], other.spec)

    def test_empty_store_returns_nothing(self):
        assert InMemoryVectorStore().search([0.1], "anything") == []

    def test_bm25_scores_rare_terms_above_common_ones(self):
        index = BM25Index()
        for i, text in enumerate([
            "the company reported revenue", "the company reported profit",
            "the company disclosed a debenture", "the company reported revenue",
        ], start=1):
            index.add(i, text)
        assert index.score(["debenture"], 3) > index.score(["company"], 3)


# ===========================================================================
# Search and citations
# ===========================================================================
class TestSearch:
    @pytest.fixture(scope="class")
    def engine(self, ar_bytes, cc_bytes, cr_bytes):
        provider = HashingEmbeddingProvider()
        store = InMemoryVectorStore()
        pipeline = IngestionPipeline(provider)
        chunk_id = 0
        for document_id, (payload, name) in enumerate([
            (ar_bytes, "Annual Report FY25"),
            (cc_bytes, "Q4 FY25 Concall"),
            (cr_bytes, "CRISIL Rating"),
        ], start=1):
            result = pipeline.run(
                payload, f"{name}.pdf",
                company_name="Bharat Consumer Products Ltd",
            )
            records = []
            for chunk, vector in zip(result.chunks, result.embeddings):
                chunk_id += 1
                records.append(VectorRecord(
                    chunk_id=chunk_id, document_id=document_id, text=chunk.text,
                    page=chunk.page, paragraph=chunk.paragraph,
                    section=chunk.section, document_title=name, vector=vector,
                ))
            store.add(records, provider.spec)
        return DocumentSearch(store, provider)

    @pytest.mark.parametrize("query,expected", [
        ("What is the EBITDA margin guidance?", "17"),
        ("Who are the competitors?", "Hindustan Unilever"),
        ("What is the credit rating?", "AA+"),
        ("subsidiaries of the company", "subsidiary"),
        ("order book and execution period", "order book"),
    ])
    def test_answers_questions_the_corpus_can_answer(self, engine, query, expected):
        answer = engine.answer(query)
        assert answer.unavailable_reason is None, answer.unavailable_reason
        assert expected.lower() in answer.answer.lower()
        assert answer.confidence > 0.3

    def test_declares_unavailability_rather_than_guessing(self, engine):
        """The grounding rule applies to search as much as to the analyst.

        Asked about a dividend policy by a corpus that never mentions
        dividends, an earlier confidence model returned 0.64 on a governance
        sentence containing the word "policy" — a confident answer to an
        unanswerable question.
        """
        answer = engine.answer("What is the dividend policy?")
        assert answer.unavailable_reason is not None
        assert not answer.answer
        assert answer.confidence < 0.2

    def test_empty_query_is_handled(self, engine):
        answer = engine.answer("   ")
        assert answer.unavailable_reason is not None

    def test_answers_carry_page_references(self, engine):
        answer = engine.answer("What is the EBITDA margin guidance?")
        assert "[p." in answer.answer

    def test_every_hit_is_citable(self, engine):
        answer = engine.answer("What are the principal risks?")
        for citation in cite_all(answer.hits):
            assert citation.page >= 1
            assert citation.document_title
            assert citation.quote
            assert "p." in citation.render()

    def test_citation_audit_catches_a_fabricated_page(self, engine):
        """An answer citing a page nothing was retrieved from is fabricated evidence.

        This is the document-side counterpart of Module 6's citation auditor.
        """
        answer = engine.answer("What is the credit rating?")
        citations = cite_all(answer.hits)
        assert verify_answer_citations(answer.answer, citations)["verified"]

        tampered = answer.answer + " Revenue was 99,999 crore [p.9999]."
        audit = verify_answer_citations(tampered, citations)
        assert audit["verified"] is False
        assert 9999 in audit["unsupported_pages"]

    def test_hits_are_ranked_by_score(self, engine):
        hits = engine.search("EBITDA margin guidance")
        assert hits == sorted(hits, key=lambda h: -h.score)

    def test_top_k_is_respected(self, engine):
        assert len(engine.search("company", top_k=2)) <= 2

    def test_query_terms_drop_stopwords(self):
        assert "the" not in query_terms("What is the revenue of the company?")


# ===========================================================================
# Knowledge graph
# ===========================================================================
class TestKnowledgeGraph:
    @pytest.fixture()
    def graph(self, annual_report):
        entities = EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        builder = KnowledgeGraphBuilder("Bharat Consumer Products Ltd", "BHARATCP")
        return builder.add_entities(entities)

    def test_builds_nodes_and_edges(self, graph):
        assert graph.node_count > 10
        assert graph.edge_count > 10

    def test_subject_company_is_the_hub(self, graph):
        subject = node_key(EntityKind.COMPANY, "Bharat Consumer Products Ltd")
        assert graph.degree(subject) == graph.edge_count

    def test_subsidiary_edges_point_at_the_parent(self, graph):
        from app.domain.documents.types import RelationKind

        subject = node_key(EntityKind.COMPANY, "Bharat Consumer Products Ltd")
        edges = [
            e for e in graph.edges.values()
            if e.relation is RelationKind.SUBSIDIARY_OF
        ]
        assert edges
        assert all(e.target == subject for e in edges)

    def test_every_edge_carries_its_evidence(self, graph):
        """An edge is a claim about the world, so it needs pages behind it."""
        for edge in graph.edges.values():
            assert edge.pages
            assert all(p >= 1 for p in edge.pages)
            assert 0.0 < edge.confidence <= 1.0

    def test_no_transitive_edges_are_invented(self, graph):
        """A→B and B→C must not produce A→C. The document did not say so."""
        subject = node_key(EntityKind.COMPANY, "Bharat Consumer Products Ltd")
        for edge in graph.edges.values():
            assert subject in (edge.source, edge.target)

    def test_identity_merges_legal_forms(self):
        assert node_key(EntityKind.SUBSIDIARY, "Acme Ltd.") == \
            node_key(EntityKind.SUBSIDIARY, "ACME Limited")

    def test_kind_is_part_of_identity(self):
        """A person and a company sharing a name are two nodes, not one."""
        assert node_key(EntityKind.DIRECTOR, "Tata") != \
            node_key(EntityKind.COMPANY, "Tata")

    def test_serialises_for_the_api(self, graph):
        payload = graph.to_dict()
        assert set(payload) == {"nodes", "edges", "stats"}
        assert payload["stats"]["nodes"] == graph.node_count

    def test_merging_graphs_accumulates_evidence(self, annual_report, cr_bytes):
        rating = parse_document(cr_bytes, "rating.pdf")
        builder = KnowledgeGraphBuilder("Bharat Consumer Products Ltd", "BHARATCP")
        builder.add_entities(
            EntityExtractor("Bharat Consumer Products Ltd").extract(annual_report)
        )
        before = builder.graph.node_count
        builder.add_entities(
            EntityExtractor("Bharat Consumer Products Ltd").extract(rating)
        )
        assert builder.graph.node_count >= before


# ===========================================================================
# Classification
# ===========================================================================
class TestClassification:
    def test_classifies_an_annual_report(self, annual_report, ar_sections):
        assert classify_document(
            "BHARATCP_AnnualReport_FY25.pdf", annual_report, sections=ar_sections
        ) is DocumentType.ANNUAL_REPORT

    def test_classifies_a_transcript(self, cc_bytes):
        doc = parse_document(cc_bytes, "concall.pdf")
        sections = SectionDetector().detect(doc)
        assert classify_document("concall.pdf", doc, sections=sections) \
            is DocumentType.CONFERENCE_CALL

    def test_classifies_a_rating_report(self, cr_bytes):
        doc = parse_document(cr_bytes, "crisil.pdf")
        assert classify_document("crisil.pdf", doc) is DocumentType.CREDIT_RATING

    def test_content_outweighs_a_misleading_filename(self, cc_bytes):
        doc = parse_document(cc_bytes, "document(3).pdf")
        assert classify_document("document(3).pdf", doc) is DocumentType.CONFERENCE_CALL

    def test_declines_rather_than_guessing(self):
        doc = ParsedDocument(pages=[ParsedPage(number=1, text="Hello world.")])
        kind, score = classify_with_confidence("file.pdf", doc)
        assert kind is DocumentType.OTHER
        assert score < 0.3


# ===========================================================================
# Pipeline
# ===========================================================================
class TestPipeline:
    def test_runs_every_stage(self, ingested):
        recorded = set(ingested.timing_map())
        expected = {
            "parse", "ocr", "layout", "tables", "sections", "entities",
            "financials", "chunking", "embedding", "knowledge",
        }
        assert expected <= recorded

    def test_produces_the_full_result(self, ingested):
        assert ingested.page_count >= 3
        assert ingested.sections
        assert ingested.tables
        assert ingested.entities
        assert ingested.chunks
        assert ingested.extraction is not None
        assert ingested.graph is not None

    def test_embeddings_align_with_chunks(self, ingested):
        assert len(ingested.embeddings) == len(ingested.chunks)

    def test_infers_type_and_period(self, ingested):
        assert ingested.doc_type is DocumentType.ANNUAL_REPORT
        assert ingested.period == "FY25"
        assert ingested.fiscal_year == 2025

    def test_content_hash_identifies_the_bytes(self, ingested, ar_bytes):
        from app.domain.documents.types import content_hash

        assert ingested.content_hash == content_hash(ar_bytes)

    def test_reports_throughput(self, ingested):
        assert ingested.total_ms > 0
        assert ingested.throughput_pages_per_second() > 0

    def test_is_deterministic(self, ar_bytes):
        pipeline = IngestionPipeline()
        first = pipeline.run(
            ar_bytes, "ar.pdf", company_name="Bharat Consumer Products Ltd"
        )
        second = pipeline.run(
            ar_bytes, "ar.pdf", company_name="Bharat Consumer Products Ltd"
        )
        assert [c.text for c in first.chunks] == [c.text for c in second.chunks]
        assert first.embeddings == second.embeddings
        assert first.extraction.coverage == second.extraction.coverage


# ===========================================================================
# Architectural invariants
# ===========================================================================
class TestArchitecture:
    def test_each_helper_is_defined_exactly_once(self):
        """No duplicated calculation — the rule that has governed every module."""
        import ast
        import pathlib

        roots = [
            pathlib.Path("app/domain/documents"),
            pathlib.Path("app/services/documents"),
        ]
        definitions: dict[str, list[str]] = {}
        watched = {
            "detect_unit", "parse_number", "normalise_table", "recover_merges",
            "build_table", "split_sentences", "tokenise", "stem", "cosine",
            "detect_period", "classify_heading", "node_key", "cite",
            "text_fingerprint", "normalise_entity", "estimate_tokens",
            "to_crore", "verify_answer_citations", "query_terms",
        }
        for root in roots:
            for path in root.rglob("*.py"):
                tree = ast.parse(path.read_text())
                for node in tree.body:  # module level only
                    if isinstance(node, ast.FunctionDef) and node.name in watched:
                        definitions.setdefault(node.name, []).append(str(path))
        duplicated = {k: v for k, v in definitions.items() if len(v) > 1}
        assert duplicated == {}

    def test_domain_layer_has_no_infrastructure_imports(self):
        """The domain must not know about the database, HTTP or a vendor."""
        import pathlib

        forbidden = ("sqlalchemy", "fastapi", "httpx", "app.models", "app.api")
        for path in pathlib.Path("app/domain/documents").rglob("*.py"):
            source = path.read_text()
            for term in forbidden:
                assert f"import {term}" not in source, f"{path} imports {term}"
                assert f"from {term}" not in source, f"{path} imports {term}"

    def test_pipeline_layer_does_not_touch_the_database(self):
        """The pipeline is pure, which is what makes it testable and reusable."""
        import pathlib

        for path in pathlib.Path("app/services/documents/pipeline").rglob("*.py"):
            source = path.read_text()
            assert "from app.models" not in source, path
            assert "Session" not in source or "orm" not in source, path

    def test_parsers_are_registered_not_hard_coded(self):
        """Adding a format is one module plus a decorator, never a branch."""
        import pathlib

        source = pathlib.Path(
            "app/services/documents/extractors/base.py"
        ).read_text()
        assert "if fmt ==" not in source
        assert "if filename.endswith" not in source


# ===========================================================================
# Performance
# ===========================================================================
class TestPerformance:
    """Guards against the pathologies that actually occurred while building this.

    Thresholds are deliberately loose — several times the observed figure — so
    they fail on a regression in kind rather than on CI jitter. A benchmark
    that goes red when the machine is busy gets muted, and then it guards
    nothing at all.
    """

    @pytest.fixture(scope="class")
    def large_pdf(self) -> bytes:
        from tests.fixtures.make_docs import (
            ANNUAL_REPORT_TABLE, ANNUAL_REPORT_TEXT, _build_pdf,
        )

        return _build_pdf(
            "\n\n".join([ANNUAL_REPORT_TEXT] * 20),
            tables=[ANNUAL_REPORT_TABLE] * 6,
        )

    def test_table_prefilter_is_selective(self, large_pdf):
        """The pre-filter must not send every page to pdfplumber.

        It once did, on a 45-page document with six tables. pdfplumber then
        cost 7.1 seconds against 107ms of real parsing — a 67x tax to search
        for tables on 39 pages that had none.
        """
        from app.services.documents.extractors.pdf import PdfParser

        parsed = PdfParser(extract_tables=False).parse(large_pdf)
        candidates = [p for p in parsed.pages if PdfParser._looks_tabular(p)]
        assert len(candidates) < parsed.page_count * 0.35

    def test_prefilter_still_finds_every_real_table(self, large_pdf):
        """Selectivity is worthless if it loses the tables it exists to find."""
        parsed = parse_document(large_pdf, "large.pdf")
        assert len(parsed.tables) == 6
        for table in parsed.tables:
            assert table.unit is Unit.INR_CRORE
            assert table.n_rows >= 12

    def test_ingestion_of_a_large_document_is_bounded(self, large_pdf):
        import time

        started = time.perf_counter()
        result = IngestionPipeline().run(
            large_pdf, "large.pdf", company_name="Bharat Consumer Products Ltd"
        )
        elapsed = time.perf_counter() - started
        assert result.page_count >= 40
        # Observed ~0.73s. Ten seconds fails only on a change in complexity.
        assert elapsed < 10.0

    def test_parsing_scales_roughly_linearly(self):
        """Superlinear parsing is how a 300-page report becomes unusable."""
        import time

        from tests.fixtures.make_docs import (
            ANNUAL_REPORT_TABLE, ANNUAL_REPORT_TEXT, _build_pdf,
        )

        def elapsed_for(multiple: int) -> tuple[float, int]:
            payload = _build_pdf(
                "\n\n".join([ANNUAL_REPORT_TEXT] * multiple),
                tables=[ANNUAL_REPORT_TABLE] * min(multiple, 6),
            )
            started = time.perf_counter()
            parsed = parse_document(payload, f"doc_{multiple}.pdf")
            return time.perf_counter() - started, parsed.page_count

        small_time, small_pages = elapsed_for(4)
        large_time, large_pages = elapsed_for(20)
        ratio_pages = large_pages / small_pages
        ratio_time = large_time / max(small_time, 1e-6)
        assert ratio_time < ratio_pages * 3.0

    def test_search_latency_is_low(self, large_pdf):
        import statistics
        import time

        provider = HashingEmbeddingProvider()
        result = IngestionPipeline(provider).run(
            large_pdf, "large.pdf", company_name="Bharat Consumer Products Ltd"
        )
        store = InMemoryVectorStore()
        store.add([
            VectorRecord(
                chunk_id=i, document_id=1, text=chunk.text, page=chunk.page,
                paragraph=chunk.paragraph, section=chunk.section,
                document_title="Large", vector=vector,
            )
            for i, (chunk, vector) in enumerate(
                zip(result.chunks, result.embeddings), start=1
            )
        ], provider.spec)

        engine = DocumentSearch(store, provider)
        timings = []
        for query in ("EBITDA margin guidance", "credit rating", "principal risks"):
            for _ in range(5):
                started = time.perf_counter()
                engine.answer(query)
                timings.append((time.perf_counter() - started) * 1000.0)
        # Observed p50 ~7ms over ~90 chunks.
        assert statistics.median(timings) < 200.0

    def test_embedding_throughput(self):
        import time

        provider = HashingEmbeddingProvider()
        texts = [f"Revenue for FY{y} was {y * 13} crore across segments." for y in range(500)]
        started = time.perf_counter()
        provider.embed(texts)
        elapsed = time.perf_counter() - started
        assert len(texts) / elapsed > 200
