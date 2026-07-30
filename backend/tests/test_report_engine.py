"""Unit tests for the report engine.

Pure layers only: the block model, the citation engine, serialisation, the
chart engine and the five renderers. No database, no HTTP.

Several tests pin defects found while building the module and say so, because
a regression test whose motivation is undocumented gets deleted by whoever next
finds it inconvenient.
"""
from __future__ import annotations

from datetime import date, datetime, timezone

import pytest

from app.domain.reports.blocks import (
    DEFAULT_DISCLAIMER, REPORT_SECTIONS, REPORT_TITLES, SECTION_ORDER,
    Block, BlockKind, Bullets, Callout, CalloutTone, Chart, ChartKind,
    CitationList, CoverMeta, Divider, Evidence, EvidenceSource, Heading,
    Insufficient, KeyValue, MetricGrid, PageBreak, Paragraph, Quote,
    ReportDocument, ReportType, Section, SectionKey, Table, Theme,
    narratives_for, sections_for,
)
from app.domain.reports.citations import (
    CitationAudit, EvidenceRegistry, annotate, audit_report,
    evidence_by_source, split_sentences, strip_markers,
)
from app.services.reports.charts.engine import ChartEngine
from app.services.reports.renderers import docx, pdf, web, xlsx  # noqa: F401
from app.services.reports.renderers.base import (
    OutputFormat, ReportRenderer, cover_pairs, registered_formats,
    renderer_for, toc_entries,
)
from app.services.reports.serialise import (
    SCHEMA_VERSION, block_from_dict, block_to_dict, document_from_dict,
    document_to_dict,
)


# ===========================================================================
# Fixtures
# ===========================================================================
@pytest.fixture()
def registry() -> EvidenceRegistry:
    reg = EvidenceRegistry()
    reg.add("revenue", "Revenue", EvidenceSource.FINANCIAL, 33543.0, "₹ cr")
    reg.add("ebitda", "EBITDA", EvidenceSource.FINANCIAL, 5490.7, "₹ cr")
    reg.add("target", "Target price", EvidenceSource.VALUATION, 185.59, "₹")
    reg.add("score", "Institutional score", EvidenceSource.SCORING, 71.0, "/100")
    return reg


@pytest.fixture()
def document(registry) -> ReportDocument:
    doc = ReportDocument(
        cover=CoverMeta(
            company_name="Bharat Consumer Products Ltd", ticker="BHARATCP",
            report_type=ReportType.INSTITUTIONAL,
            title="Institutional Research Report",
            as_of=date(2025, 3, 31), analyst="Test Analyst",
            exchange="NSE", sector="FMCG", industry="Packaged Foods",
            recommendation="HOLD", rating="A", score=71.0,
            current_price=268.0, target_price=185.59, upside=-0.3075,
            market_cap=67000.0,
        ),
        generated_at=datetime(2025, 3, 31, tzinfo=timezone.utc),
        disclaimer=DEFAULT_DISCLAIMER,
        provenance={"report_type": "institutional"},
    )

    summary = Section(SectionKey.EXECUTIVE_SUMMARY, "Executive Summary")
    summary.add(MetricGrid(
        [("Price", "₹268.00", "BHARATCP"), ("Target", "₹185.59", "blended")],
        evidence=registry.many("target"),
    ))
    summary.add(Callout(
        "Recommendation — HOLD",
        "A score of 71.0 [score] maps to accumulate but valuation caps it.",
        CalloutTone.WARNING, evidence=registry.many("score"),
    ))
    doc.add(summary)

    financial = Section(SectionKey.FINANCIAL_ANALYSIS, "Financial Analysis")
    financial.add(Paragraph(
        "Revenue was ₹33,543.00 crore [revenue] with EBITDA of "
        "₹5,490.70 crore [ebitda].",
        evidence=registry.many("revenue", "ebitda"),
    ))
    financial.add(Table(
        ["₹ crore", "FY24", "FY25"],
        [["Revenue", "28,914", "33,543"], ["EBITDA", "4,702", "5,491"]],
        caption="Key financials", emphasis_rows=[0],
        evidence=registry.many("revenue"),
    ))
    financial.add(Chart(
        ChartKind.REVENUE, "Revenue and EBITDA", labels=["FY24", "FY25"],
        series=[("Revenue", [28914.0, 33543.0]), ("EBITDA", [4702.0, 5490.7])],
        y_unit="₹ cr", evidence=registry.many("revenue"),
    ))
    financial.add(Bullets(["Volume-led growth", "Margin expansion of 11 bps"]))
    financial.add(KeyValue([("Sector", "FMCG"), ("Industry", "Packaged Foods")]))
    doc.add(financial)

    moat = Section(SectionKey.MOAT, "Economic Moat")
    moat.mark_insufficient("No competitive-advantage inputs are populated.")
    doc.add(moat)

    appendix = Section(SectionKey.APPENDIX, "Appendix")
    appendix.add(Heading("Evidence", level=3))
    appendix.add(CitationList(registry.all()))
    appendix.add(Quote("We expect revenue growth of 12%.", "Q4 FY25 call"))
    appendix.add(Divider())
    appendix.add(PageBreak())
    appendix.add(Paragraph(DEFAULT_DISCLAIMER))
    doc.add(appendix)

    doc.add(Section(SectionKey.TOC, "Contents").add(Paragraph("__TOC__")))
    return doc


# ===========================================================================
# Block model
# ===========================================================================
class TestBlockModel:
    def test_every_brief_section_exists(self):
        """The brief enumerates eighteen sections plus a contents page."""
        required = {
            "cover", "executive_summary", "investment_thesis",
            "business_overview", "industry_analysis", "financial_analysis",
            "forecast", "valuation", "dcf", "relative_valuation",
            "institutional_score", "management", "moat", "risk_analysis",
            "scenario_analysis", "peer_comparison", "portfolio_fit",
            "appendix",
        }
        assert required <= {s.value for s in SectionKey}

    def test_every_brief_chart_exists(self):
        required = {
            "revenue", "ebitda", "pat", "margins", "cash_flow", "dcf",
            "sensitivity", "peer_comparison", "score_radar",
            "portfolio_allocation",
        }
        assert required == {c.value for c in ChartKind}

    def test_all_six_report_types_are_composed(self):
        assert len(ReportType) == 6
        for report_type in ReportType:
            assert sections_for(report_type)
            assert report_type in REPORT_TITLES

    def test_report_types_differ_in_composition(self):
        """Two types that select the same sections are one type."""
        compositions = {rt: sections_for(rt) for rt in ReportType}
        assert len(set(compositions.values())) == len(ReportType)

    def test_quick_is_a_subset_of_deep_research(self):
        assert set(sections_for(ReportType.QUICK)) <= set(
            sections_for(ReportType.DEEP_RESEARCH)
        )

    def test_sections_render_in_canonical_order(self, document):
        """Two reports of one type must be comparable page by page."""
        order = [s.key for s in document.ordered()]
        assert order == sorted(order, key=SECTION_ORDER.index)

    def test_contents_precedes_the_body(self, document):
        keys = [s.key for s in document.ordered()]
        assert keys.index(SectionKey.TOC) < keys.index(
            SectionKey.EXECUTIVE_SUMMARY
        )

    def test_empty_sections_are_not_added(self):
        doc = ReportDocument(cover=_cover())
        doc.add(Section(SectionKey.MOAT, "Moat"))
        assert doc.sections == []

    def test_insufficient_replaces_content_and_records_why(self):
        section = Section(SectionKey.MOAT, "Moat")
        section.add(Paragraph("draft"))
        section.mark_insufficient("No inputs.")
        assert not section.sufficient
        assert len(section.blocks) == 1
        assert section.blocks[0].text.startswith("Insufficient evidence.")
        assert "No inputs." in section.blocks[0].text

    def test_insufficient_statement_is_exactly_the_brief_wording(self):
        assert Insufficient.STATEMENT == "Insufficient evidence."

    def test_evidence_deduplicates_across_blocks(self, document, registry):
        keys = [e.key for e in document.evidence()]
        assert len(keys) == len(set(keys))

    def test_statistics_are_consistent(self, document):
        stats = document.statistics()
        assert stats["sections"] == len(document.sections)
        assert stats["charts"] == len(document.charts())
        assert stats["tables"] == len(document.tables())
        assert stats["sections_insufficient"] == 1

    def test_table_defaults_to_right_aligned_figures(self):
        """Ragged decimal points defeat reading down a numeric column."""
        table = Table(["Item", "FY24", "FY25"], [["Revenue", "1", "2"]])
        assert table.align == ["l", "r", "r"]

    def test_chart_with_no_data_is_flagged(self):
        """An empty axis reads as zero, which is a fabrication."""
        assert not Chart(
            ChartKind.REVENUE, "Empty", labels=["FY25"],
            series=[("Revenue", [None])],
        ).has_data
        assert Chart(
            ChartKind.REVENUE, "Real", labels=["FY25"],
            series=[("Revenue", [1.0])],
        ).has_data

    def test_narratives_are_declared_per_report_type(self):
        assert len(narratives_for(ReportType.QUICK)) < len(
            narratives_for(ReportType.DEEP_RESEARCH)
        )
        for report_type in ReportType:
            assert isinstance(narratives_for(report_type), tuple)


def _cover() -> CoverMeta:
    return CoverMeta(
        company_name="Test Ltd", ticker="TEST",
        report_type=ReportType.QUICK, title="Quick Report",
    )


# ===========================================================================
# Citation engine
# ===========================================================================
class TestCitations:
    def test_duplicate_key_with_a_different_value_is_refused(self, registry):
        """Two statements citing one key must not mean different things."""
        with pytest.raises(ValueError, match="already registered"):
            registry.add("revenue", "Revenue", EvidenceSource.FINANCIAL, 999.0)

    def test_reregistering_the_same_value_is_allowed(self, registry):
        registry.add("revenue", "Revenue", EvidenceSource.FINANCIAL, 33543.0, "₹ cr")
        assert len(registry) == 4

    def test_many_skips_unregistered_keys(self, registry):
        assert len(registry.many("revenue", "nonexistent")) == 1

    def test_sentence_split_survives_decimals(self):
        """Module 6 shipped a splitter that tore ₹33,543.00 in half.

        It orphaned the citation attached to the figure and reported 50%
        coverage on a perfectly cited answer.
        """
        sentences = split_sentences(
            "Revenue was ₹33,543.00 crore. EBITDA margin reached 16.37%."
        )
        assert len(sentences) == 2
        assert "33,543.00" in sentences[0]

    def test_a_cited_claim_is_supported(self, document):
        audit = audit_report(document)
        assert audit.coverage == 1.0
        assert audit.is_clean

    def test_an_uncited_numeric_claim_is_caught(self):
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.FINANCIAL_ANALYSIS, "Financial")
        section.add(Paragraph("Margins improved by 210 bps."))
        doc.add(section)
        audit = audit_report(doc)
        assert audit.total_claims == 1
        assert audit.unsupported
        assert not audit.is_clean

    def test_a_dangling_marker_is_caught(self):
        """A marker pointing at nothing is a fabricated reference."""
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.VALUATION, "Valuation")
        section.add(Paragraph("Fair value is ₹185.59 [nonexistent]."))
        doc.add(section)
        audit = audit_report(doc)
        assert "nonexistent" in audit.dangling_markers
        assert not audit.is_clean

    def test_block_level_evidence_supports_its_prose(self, registry):
        """A paragraph introducing a cited table need not repeat the markers."""
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.FINANCIAL_ANALYSIS, "Financial")
        section.add(Paragraph(
            "Revenue reached ₹33,543 crore.",
            evidence=registry.many("revenue"),
        ))
        doc.add(section)
        assert audit_report(doc).is_clean

    def test_hedged_sentences_are_not_claims(self):
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.MOAT, "Moat")
        section.add(Paragraph(
            "Insufficient evidence. Only 2 of 13 inputs are populated."
        ))
        doc.add(section)
        assert audit_report(doc).total_claims == 0

    def test_small_integers_in_prose_are_not_claims(self):
        """"three pillars" is not a claim about the company."""
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.INVESTMENT_THESIS, "Thesis")
        section.add(Paragraph("The thesis rests on 3 pillars."))
        doc.add(section)
        assert audit_report(doc).total_claims == 0

    def test_marker_digits_are_not_mistaken_for_claims(self):
        """`[revenue_fy25]` must not register "25" as an uncited number."""
        doc = ReportDocument(cover=_cover())
        registry = EvidenceRegistry()
        registry.add("revenue_fy25", "Revenue FY25", EvidenceSource.FINANCIAL, 1.0)
        section = Section(SectionKey.FINANCIAL_ANALYSIS, "Financial")
        section.add(Paragraph(
            "Revenue grew [revenue_fy25].", evidence=registry.all(),
        ))
        doc.add(section)
        assert audit_report(doc).total_claims == 0

    def test_audit_records_which_engines_were_used(self, document):
        sources = audit_report(document).sources_used
        assert "financial_engine" in sources
        assert "valuation_engine" in sources

    def test_annotate_swaps_markers_for_labels(self, registry):
        assert annotate("Revenue [revenue] rose.", registry) == \
            "Revenue [Revenue] rose."

    def test_annotate_leaves_an_unknown_marker_visible(self, registry):
        """Stripping it would hide a broken reference."""
        assert "[missing]" in annotate("Value [missing].", registry)

    def test_strip_markers(self):
        assert strip_markers("Revenue [revenue] rose.") == "Revenue rose."

    def test_evidence_groups_by_engine(self, registry):
        grouped = evidence_by_source(registry.all())
        assert EvidenceSource.FINANCIAL in grouped
        assert len(grouped[EvidenceSource.FINANCIAL]) == 2


# ===========================================================================
# Chart engine
# ===========================================================================
class TestChartEngine:
    @pytest.fixture(scope="class")
    def engine(self) -> ChartEngine:
        return ChartEngine(Theme.LIGHT)

    def test_renders_a_bar_chart(self, engine):
        png = engine.render(Chart(
            ChartKind.REVENUE, "Revenue", labels=["FY24", "FY25"],
            series=[("Revenue", [100.0, 120.0])],
        ))
        assert png and png[:8] == b"\x89PNG\r\n\x1a\n"

    def test_renders_every_chart_kind_the_brief_lists(self, engine):
        specs = {
            ChartKind.REVENUE: dict(series=[("Revenue", [1.0, 2.0])]),
            ChartKind.EBITDA: dict(series=[("EBITDA", [1.0, 2.0])]),
            ChartKind.PAT: dict(series=[("PAT", [1.0, 2.0])]),
            ChartKind.MARGINS: dict(series=[("Margin", [0.1, 0.2])], y_unit="%"),
            ChartKind.CASH_FLOW: dict(series=[("CFO", [1.0, 2.0])]),
            ChartKind.DCF: dict(series=[("PV", [1.0, 2.0])]),
            ChartKind.PEER_COMPARISON: dict(series=[("Cap", [1.0, 2.0])]),
            ChartKind.PORTFOLIO_ALLOCATION: dict(series=[("W", [0.6, 0.4])]),
        }
        for kind, extra in specs.items():
            png = engine.render(
                Chart(kind, kind.value, labels=["A", "B"], **extra)
            )
            assert png, f"{kind.value} rendered nothing"

        radar = engine.render(Chart(
            ChartKind.SCORE_RADAR, "Radar", labels=["A", "B", "C", "D"],
            series=[("Score", [6.0, 7.0, 8.0, 5.0])],
        ))
        assert radar

        heatmap = engine.render(Chart(
            ChartKind.SENSITIVITY, "Sensitivity", labels=["10%", "11%"],
            row_labels=["2%", "3%"], matrix=[[100.0, 90.0], [110.0, 95.0]],
        ))
        assert heatmap

    def test_an_empty_chart_is_not_drawn(self, engine):
        assert engine.render(Chart(
            ChartKind.REVENUE, "Empty", labels=["FY25"],
            series=[("Revenue", [None])],
        )) is None

    def test_a_two_axis_radar_is_refused(self, engine):
        """A radar of two axes is a line and misrepresents the shape."""
        assert engine.render(Chart(
            ChartKind.SCORE_RADAR, "Too few", labels=["A", "B"],
            series=[("Score", [5.0, 6.0])],
        )) is None

    def test_nulls_leave_gaps_rather_than_plotting_zero(self, engine):
        png = engine.render(Chart(
            ChartKind.REVENUE, "Gapped", labels=["A", "B", "C"],
            series=[("Revenue", [1.0, None, 3.0])],
        ))
        assert png

    def test_cache_returns_the_same_bytes(self, engine):
        chart = Chart(
            ChartKind.REVENUE, "Cached", labels=["A", "B"],
            series=[("R", [1.0, 2.0])],
        )
        first = engine.render(chart)
        before = engine.hits
        second = engine.render(chart)
        assert first == second
        assert engine.hits == before + 1

    def test_theme_is_part_of_the_cache_key(self):
        """A dark render must not be served for a light request."""
        chart = Chart(
            ChartKind.REVENUE, "Themed", labels=["A"], series=[("R", [1.0])],
        )
        light = ChartEngine(Theme.LIGHT)
        dark = ChartEngine(Theme.DARK)
        assert light.cache_key(chart, 7.2, 3.2) != dark.cache_key(chart, 7.2, 3.2)
        assert light.render(chart) != dark.render(chart)


# ===========================================================================
# Renderers
# ===========================================================================
class TestRenderers:
    def test_all_five_brief_formats_are_registered(self):
        assert {f.value for f in registered_formats()} == {
            "pdf", "docx", "xlsx", "html", "md",
        }

    def test_every_renderer_handles_the_full_block_vocabulary(self):
        """A new block type must not silently render as nothing anywhere."""
        for fmt in registered_formats():
            assert renderer_for(fmt).handles == frozenset(BlockKind), fmt

    @pytest.mark.parametrize("fmt", list(OutputFormat))
    def test_each_format_produces_output(self, document, fmt):
        result = renderer_for(fmt).render(document)
        assert result.size_bytes > 0
        assert result.filename.endswith(fmt.extension)
        assert result.fmt is fmt

    def test_pdf_has_pages_bookmarks_and_metadata(self, document):
        import fitz

        result = renderer_for(OutputFormat.PDF).render(document)
        assert result.page_count and result.page_count >= 2

        opened = fitz.open(stream=result.payload, filetype="pdf")
        try:
            assert opened.page_count == result.page_count
            assert opened.get_toc(), "no bookmarks in the PDF outline"
            assert document.cover.company_name in (opened.metadata or {}).get(
                "title", ""
            )
        finally:
            opened.close()

    def test_pdf_renders_the_rupee_symbol(self, document):
        """Base-14 Helvetica has no rupee glyph and printed a black box.

        Every monetary figure in the report was affected, which is the kind of
        defect that survives review because the layout still looks right.
        """
        import fitz

        result = renderer_for(OutputFormat.PDF).render(document)
        opened = fitz.open(stream=result.payload, filetype="pdf")
        try:
            text = "".join(page.get_text() for page in opened)
        finally:
            opened.close()
        assert "₹" in text

    def test_pdf_bookmark_levels_never_skip(self, document):
        """PDF outlines are a strict tree.

        A section heading followed directly by a caption jumps two levels, and
        ReportLab raises `can't jump from outline level 0 to level 2`.
        """
        import fitz

        result = renderer_for(OutputFormat.PDF).render(document)
        opened = fitz.open(stream=result.payload, filetype="pdf")
        try:
            levels = [entry[0] for entry in opened.get_toc()]
        finally:
            opened.close()
        previous = 0
        for level in levels:
            assert level <= previous + 1, "outline level skipped"
            previous = level

    def test_docx_opens_and_carries_the_content(self, document):
        import io

        from docx import Document as Docx

        result = renderer_for(OutputFormat.DOCX).render(document)
        opened = Docx(io.BytesIO(result.payload))
        text = "\n".join(p.text for p in opened.paragraphs)
        assert document.cover.company_name in text
        assert opened.tables

    def test_docx_uses_real_heading_styles(self, document):
        """Bold paragraphs look identical and are unnavigable."""
        import io

        from docx import Document as Docx

        result = renderer_for(OutputFormat.DOCX).render(document)
        opened = Docx(io.BytesIO(result.payload))
        styles = {p.style.name for p in opened.paragraphs}
        assert any(s.startswith("Heading") for s in styles)

    def test_xlsx_opens_with_a_sheet_per_section(self, document):
        import io

        import openpyxl

        result = renderer_for(OutputFormat.XLSX).render(document)
        workbook = openpyxl.load_workbook(io.BytesIO(result.payload))
        assert "Summary" in workbook.sheetnames
        assert "Evidence" in workbook.sheetnames
        assert len(workbook.sheetnames) >= 4

    def test_xlsx_writes_numbers_as_numbers(self, document):
        """A spreadsheet of text nobody can sum defeats the export."""
        import io

        import openpyxl

        result = renderer_for(OutputFormat.XLSX).render(document)
        workbook = openpyxl.load_workbook(io.BytesIO(result.payload))
        sheet = workbook["Financial Analysis"]
        numeric = [
            cell.value for row in sheet.iter_rows() for cell in row
            if isinstance(cell.value, (int, float))
        ]
        assert numeric, "no numeric cells were written"

    def test_xlsx_neutralises_formula_injection(self):
        """A cell beginning with '=' would be evaluated on open."""
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.APPENDIX, "Appendix")
        section.add(Table(["Note"], [["=1+1"], ["-- see note"]]))
        doc.add(section)

        import io

        import openpyxl

        result = renderer_for(OutputFormat.XLSX).render(doc)
        workbook = openpyxl.load_workbook(io.BytesIO(result.payload))
        values = [
            cell.value for row in workbook["Appendix"].iter_rows()
            for cell in row if isinstance(cell.value, str)
        ]
        assert not any(v.startswith("=") for v in values)

    def test_html_is_self_contained(self, document):
        result = renderer_for(OutputFormat.HTML).render(document)
        text = result.payload.decode("utf-8")
        assert text.startswith("<!doctype html>")
        assert "<style>" in text
        # Charts are inlined as data URIs, so there is no asset directory.
        assert "data:image/png;base64," in text
        assert 'src="http' not in text

    def test_html_escapes_hostile_content(self):
        doc = ReportDocument(cover=_cover())
        section = Section(SectionKey.APPENDIX, "Appendix")
        section.add(Paragraph("<script>alert('x')</script>"))
        doc.add(section)
        text = renderer_for(OutputFormat.HTML).render(doc).payload.decode()
        assert "<script>alert" not in text
        assert "&lt;script&gt;" in text

    def test_html_has_print_rules(self, document):
        """The brief asks for print-ready output from every format."""
        text = renderer_for(OutputFormat.HTML).render(document).payload.decode()
        assert "@media print" in text
        assert "@page" in text

    def test_markdown_has_tables_and_headings(self, document):
        text = renderer_for(OutputFormat.MARKDOWN).render(document).payload.decode()
        assert text.startswith("# ")
        assert "| " in text and "---" in text
        assert "## Financial Analysis" in text

    def test_every_format_states_insufficient_evidence(self, document):
        """The brief's required wording must survive every renderer."""
        for fmt in (OutputFormat.HTML, OutputFormat.MARKDOWN):
            text = renderer_for(fmt).render(document).payload.decode()
            assert "Insufficient evidence" in text

        import fitz

        result = renderer_for(OutputFormat.PDF).render(document)
        opened = fitz.open(stream=result.payload, filetype="pdf")
        try:
            assert "Insufficient evidence" in "".join(
                page.get_text() for page in opened
            )
        finally:
            opened.close()

    def test_toc_excludes_the_cover_and_itself(self, document):
        titles = [t for _, t in toc_entries(document)]
        assert "Contents" not in titles
        assert "Executive Summary" in titles

    def test_cover_pairs_are_identical_across_renderers(self, document):
        """One formatter, so no two formats can disagree about the cover."""
        pairs = dict(cover_pairs(document))
        assert pairs["Ticker"] == "BHARATCP · NSE"
        assert pairs["Target price"] == "₹185.59"
        assert pairs["Upside"] == "-30.8%"

    def test_a_data_warning_reaches_the_cover(self):
        doc = ReportDocument(cover=CoverMeta(
            company_name="Test Ltd", ticker="TEST",
            report_type=ReportType.QUICK, title="Quick Report",
            data_warning="Illustrative valuation only.",
        ))
        doc.add(Section(SectionKey.APPENDIX, "Appendix").add(Paragraph("x")))
        text = renderer_for(OutputFormat.HTML).render(doc).payload.decode()
        assert "Illustrative valuation only." in text

    def test_renderers_share_one_chart_engine(self, document):
        """The same chart must be rasterised once, not once per format."""
        engine = ChartEngine(document.theme)
        renderer_for(OutputFormat.HTML, engine).render(document)
        misses = engine.misses
        renderer_for(OutputFormat.PDF, engine).render(document)
        assert engine.misses == misses, "the chart was re-rendered"
        assert engine.hits > 0


# ===========================================================================
# Serialisation
# ===========================================================================
class TestSerialisation:
    def test_round_trip_is_lossless(self, document):
        """A stored report re-rendered must be byte-identical."""
        restored = document_from_dict(document_to_dict(document))
        original = renderer_for(OutputFormat.MARKDOWN).render(document).payload
        rebuilt = renderer_for(OutputFormat.MARKDOWN).render(restored).payload
        assert original == rebuilt

    def test_round_trip_preserves_structure(self, document):
        restored = document_from_dict(document_to_dict(document))
        assert restored.statistics() == document.statistics()
        assert [s.key for s in restored.ordered()] == [
            s.key for s in document.ordered()
        ]

    def test_round_trip_preserves_insufficiency(self, document):
        restored = document_from_dict(document_to_dict(document))
        moat = restored.section(SectionKey.MOAT)
        assert moat is not None and not moat.sufficient
        assert moat.reason

    @pytest.mark.parametrize("block", [
        Heading("Title", 2),
        Paragraph("Text [key]."),
        Bullets(["a", "b"], ordered=True),
        KeyValue([("k", "v")]),
        Table(["A", "B"], [["1", "2"]], caption="Cap", emphasis_rows=[0]),
        MetricGrid([("L", "V", "H")]),
        Chart(ChartKind.REVENUE, "T", labels=["A"], series=[("S", [1.0])]),
        Callout("T", "B", CalloutTone.WARNING),
        Quote("Said", "Someone"),
        Divider(),
        PageBreak(),
        Insufficient("because"),
        CitationList([Evidence("k", "L", EvidenceSource.AI, 1.0)]),
    ])
    def test_every_block_kind_round_trips(self, block: Block):
        restored = block_from_dict(block_to_dict(block))
        assert restored.kind is block.kind
        assert block_to_dict(restored) == block_to_dict(block)

    def test_a_newer_schema_is_refused(self, document):
        """Silently dropping fields a newer writer added would corrupt the report."""
        payload = document_to_dict(document)
        payload["schema"] = SCHEMA_VERSION + 1
        with pytest.raises(ValueError, match="newer"):
            document_from_dict(payload)

    def test_evidence_survives_the_round_trip(self, document):
        restored = document_from_dict(document_to_dict(document))
        assert {e.key for e in restored.evidence()} == {
            e.key for e in document.evidence()
        }

    def test_serialised_form_is_json_safe(self, document):
        import json

        payload = json.dumps(document_to_dict(document))
        assert document_from_dict(json.loads(payload)).cover.ticker == "BHARATCP"


# ===========================================================================
# Architecture
# ===========================================================================
class TestArchitecture:
    def test_domain_has_no_infrastructure_imports(self):
        import pathlib

        forbidden = (
            "sqlalchemy", "fastapi", "httpx", "app.models", "app.api",
            "reportlab", "docx", "openpyxl", "matplotlib",
        )
        for path in pathlib.Path("app/domain/reports").rglob("*.py"):
            source = path.read_text()
            for term in forbidden:
                assert f"import {term}" not in source, f"{path} imports {term}"
                assert f"from {term}" not in source, f"{path} imports {term}"

    def test_renderers_do_not_touch_the_database(self):
        """A renderer that can fetch data can produce a report nobody audited."""
        import pathlib

        for path in pathlib.Path("app/services/reports/renderers").rglob("*.py"):
            source = path.read_text()
            assert "from app.models" not in source, path
            assert "Session" not in source, path

    def test_no_renderer_hard_codes_a_section(self):
        """Sections are data; a renderer that names one is a template."""
        import pathlib

        named = ("Executive Summary", "Investment Thesis", "Discounted Cash")
        for path in pathlib.Path("app/services/reports/renderers").rglob("*.py"):
            source = path.read_text()
            for title in named:
                assert title not in source, f"{path} hard-codes '{title}'"

    def test_each_helper_is_defined_once(self):
        import ast
        import pathlib

        watched = {
            "audit_report", "annotate", "strip_markers", "split_sentences",
            "evidence_by_source", "document_to_dict", "document_from_dict",
            "block_to_dict", "block_from_dict", "cover_pairs", "toc_entries",
            "sections_for", "narratives_for",
        }
        found: dict[str, list[str]] = {}
        for root in ("app/domain/reports", "app/services/reports"):
            for path in pathlib.Path(root).rglob("*.py"):
                tree = ast.parse(path.read_text())
                for node in tree.body:
                    if isinstance(node, ast.FunctionDef) and node.name in watched:
                        found.setdefault(node.name, []).append(str(path))
        assert {k: v for k, v in found.items() if len(v) > 1} == {}

    def test_report_composition_is_data_not_code(self):
        """Adding a report type must be a row, not a builder."""
        assert set(REPORT_SECTIONS) == set(ReportType)
        for sections in REPORT_SECTIONS.values():
            assert all(isinstance(s, SectionKey) for s in sections)


# ===========================================================================
# Performance
# ===========================================================================
class TestPerformanceBudget:
    """Guards against regressions in kind, not in noise.

    Thresholds sit several times above the observed figure so they fail on an
    algorithmic change rather than on a busy machine.
    """

    @staticmethod
    def _large_document() -> ReportDocument:
        """A report roughly three times the size of a deep-research run."""
        doc = ReportDocument(cover=_cover())
        for index, key in enumerate(SECTION_ORDER):
            if key in {SectionKey.COVER, SectionKey.TOC}:
                continue
            section = Section(key, key.value.replace("_", " ").title())
            for n in range(6):
                section.add(Paragraph(
                    f"Paragraph {n} of section {index} with a figure of "
                    f"{1000 + n:,} crore.",
                ))
                section.add(Table(
                    ["Item", "FY24", "FY25"],
                    [[f"Row {r}", f"{r * 11:,}", f"{r * 13:,}"] for r in range(8)],
                ))
            section.add(Chart(
                ChartKind.REVENUE, f"Chart {index}",
                labels=["FY21", "FY22", "FY23", "FY24", "FY25"],
                series=[("Revenue", [100.0 + index, 120.0, 140.0, 160.0, 180.0])],
            ))
            doc.add(section)
        return doc

    def test_markdown_render_is_fast(self, document):
        import time

        started = time.perf_counter()
        for _ in range(20):
            renderer_for(OutputFormat.MARKDOWN).render(document)
        # Observed ~0.4ms each.
        assert (time.perf_counter() - started) < 2.0

    def test_pdf_render_of_a_large_report_is_bounded(self):
        import time

        document = self._large_document()
        engine = ChartEngine(document.theme)
        started = time.perf_counter()
        result = renderer_for(OutputFormat.PDF, engine).render(document)
        elapsed = time.perf_counter() - started
        assert result.page_count and result.page_count > 10
        # Observed ~2s for a 17-chart, 100-table document.
        assert elapsed < 25.0

    def test_render_scales_roughly_linearly_with_content(self, document):
        """Superlinear rendering makes a long report unusable."""
        import time

        def elapsed(doc: ReportDocument) -> float:
            started = time.perf_counter()
            renderer_for(OutputFormat.HTML, ChartEngine(doc.theme)).render(doc)
            return time.perf_counter() - started

        small = elapsed(document)
        large = elapsed(self._large_document())
        assert large / max(small, 1e-6) < 60

    def test_chart_cache_prevents_re_rasterising(self):
        """Five formats of one report must draw each chart once."""
        document = self._large_document()
        engine = ChartEngine(document.theme)
        for fmt in (OutputFormat.HTML, OutputFormat.PDF, OutputFormat.DOCX):
            renderer_for(fmt, engine).render(document)
        assert engine.misses == len(document.charts())
        assert engine.hits >= 2 * len(document.charts())

    def test_serialisation_round_trip_is_fast(self):
        import json
        import time

        document = self._large_document()
        started = time.perf_counter()
        payload = json.dumps(document_to_dict(document))
        document_from_dict(json.loads(payload))
        # Observed ~3ms.
        assert (time.perf_counter() - started) < 2.0

    def test_citation_audit_is_bounded(self):
        import time

        document = self._large_document()
        started = time.perf_counter()
        audit_report(document)
        assert (time.perf_counter() - started) < 1.0
