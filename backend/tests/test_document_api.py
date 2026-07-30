"""Integration tests for the document-intelligence API.

These exercise the real FastAPI app against the shared seeded database from
``conftest.py``, uploading genuine PDFs through the HTTP surface.

Uploads happen once, at module scope, because ingesting three PDFs per test
would dominate the suite's runtime for no additional coverage.
"""
from __future__ import annotations

import pytest

from app.domain.documents.fields import FIELD_COUNT
from tests.fixtures.make_docs import (
    annual_report_pdf, concall_pdf, credit_rating_pdf, scanned_pdf,
)

BASE = "/api/v1"


@pytest.fixture(scope="module")
def company_id(api_client) -> str:
    response = api_client.get(f"{BASE}/companies", params={"page_size": 60})
    assert response.status_code == 200
    results = response.json()["results"]
    return next(c["id"] for c in results if c["ticker"] == "BHARATCP")


@pytest.fixture(scope="module")
def corpus_bytes() -> dict[str, bytes]:
    """Render each fixture PDF exactly once.

    PyMuPDF stamps a creation timestamp into every file it writes, so calling
    `annual_report_pdf()` twice yields different bytes. Any test about byte
    identity must therefore compare the *same* payload, not a freshly rendered
    one — the first version of the duplicate test re-rendered and then blamed
    the product for correctly noticing the bytes had changed.
    """
    return {
        "annual_report": annual_report_pdf(),
        "concall": concall_pdf(),
        "rating": credit_rating_pdf(),
    }


@pytest.fixture(scope="module")
def uploaded(api_client, company_id, corpus_bytes) -> dict[str, int]:
    """Ingest the corpus once. Returns a name → document id map."""
    ids: dict[str, int] = {}
    corpus = [
        ("annual_report", "BHARATCP_AnnualReport_FY25.pdf",
         corpus_bytes["annual_report"]),
        ("concall", "BHARATCP_Q4FY25_concall_transcript.pdf",
         corpus_bytes["concall"]),
        ("rating", "CRISIL_rating_BHARATCP.pdf", corpus_bytes["rating"]),
    ]
    for key, filename, payload in corpus:
        response = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={"file": (filename, payload, "application/pdf")},
        )
        assert response.status_code == 201, response.text
        ids[key] = response.json()["document"]["id"]
    return ids


# ===========================================================================
# Upload
# ===========================================================================
class TestUpload:
    def test_ingests_an_annual_report(self, api_client, uploaded):
        response = api_client.get(f"{BASE}/documents/{uploaded['annual_report']}")
        assert response.status_code == 200
        document = response.json()
        assert document["status"] == "ready"
        assert document["doc_type"] == "annual_report"
        assert document["period"] == "FY25"
        assert document["fiscal_year"] == 2025
        assert document["page_count"] >= 3
        assert document["chunk_count"] > 0
        assert document["table_count"] >= 1
        assert document["fact_count"] > 20
        assert document["coverage"] > 0.4

    def test_classifies_a_transcript_and_a_rating(self, api_client, uploaded):
        concall = api_client.get(f"{BASE}/documents/{uploaded['concall']}").json()
        rating = api_client.get(f"{BASE}/documents/{uploaded['rating']}").json()
        assert concall["doc_type"] == "conference_call"
        assert rating["doc_type"] == "credit_rating"

    def test_identical_bytes_are_recognised_as_a_duplicate(
        self, api_client, company_id, uploaded, corpus_bytes
    ):
        """Re-uploading the same file must not re-index or duplicate it."""
        response = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={
                "file": (
                    "BHARATCP_AnnualReport_FY25.pdf",
                    corpus_bytes["annual_report"],
                    "application/pdf",
                )
            },
        )
        assert response.status_code == 201
        payload = response.json()
        assert payload["action"] == "duplicate"
        assert payload["duplicate_of"] == uploaded["annual_report"]

    def test_same_name_different_bytes_creates_a_version(
        self, api_client, company_id, corpus_bytes
    ):
        """The predecessor is superseded, never deleted.

        A citation issued last quarter must still resolve to the text it
        actually quoted, so the old version stays in the database and only
        leaves the search index.
        """
        # Distinct content, or the hash-based deduplication correctly fires
        # first: the same bytes under a new name are still the same document,
        # which is the behaviour the previous test asserts.
        name = "BHARATCP_Versioned.pdf"
        original = corpus_bytes["annual_report"] + b"\n%v1\n"
        first = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={"file": (name, original, "application/pdf")},
        ).json()
        assert first["action"] == "created"

        second = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={"file": (name, original + b"%v2 revised\n", "application/pdf")},
        ).json()
        assert second["action"] == "new_version"
        assert second["superseded"] == first["document"]["id"]
        assert second["document"]["version"] == 2

        superseded = api_client.get(
            f"{BASE}/documents/{first['document']['id']}"
        ).json()
        assert superseded["superseded_by"] == second["document"]["id"]

    def test_unsupported_format_is_refused(self, api_client, company_id):
        response = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={"file": ("notes.rtf", b"{\\rtf1}", "application/rtf")},
        )
        assert response.status_code == 415

    def test_empty_file_is_refused(self, api_client, company_id):
        response = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={"file": ("empty.txt", b"", "text/plain")},
        )
        assert response.status_code == 400

    def test_unknown_company_is_refused(self, api_client):
        response = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": "does-not-exist"},
            files={"file": ("note.txt", b"hello world", "text/plain")},
        )
        assert response.status_code == 400

    def test_ingests_every_supported_format(self, api_client, company_id):
        import io

        import docx
        import openpyxl

        buffer = io.BytesIO()
        document = docx.Document()
        document.add_heading("Business Overview", level=1)
        document.add_paragraph("The Company is engaged in packaged foods.")
        document.save(buffer)

        workbook_buffer = io.BytesIO()
        workbook = openpyxl.Workbook()
        workbook.active.append(["Particulars", "FY25"])
        workbook.active.append(["Revenue", 33543])
        workbook.save(workbook_buffer)

        cases = [
            ("notes.txt", b"Revenue from operations grew to Rs 33,543 crore in FY25."),
            ("notes.md", b"# Overview\n\nEBITDA margin guidance of 17%.\n"),
            ("filing.html", b"<html><body><h1>Outcome</h1><p>Capex Rs 1,200 crore.</p></body></html>"),
            ("figures.csv", b"Particulars,FY25\nRevenue (Rs cr),33543\n"),
            ("model.xlsx", workbook_buffer.getvalue()),
            ("overview.docx", buffer.getvalue()),
        ]
        for filename, payload in cases:
            response = api_client.post(
                f"{BASE}/documents/upload",
                data={"company_id": company_id},
                files={"file": (filename, payload, "application/octet-stream")},
            )
            assert response.status_code == 201, f"{filename}: {response.text}"
            assert response.json()["document"]["status"] == "ready", filename


# ===========================================================================
# Listing and detail
# ===========================================================================
class TestListing:
    def test_lists_documents_for_a_company(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents", params={"company_id": company_id}
        )
        assert response.status_code == 200
        documents = response.json()
        assert len(documents) >= 3
        assert all(d["company_id"] == company_id for d in documents)

    def test_filters_by_type(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents",
            params={"company_id": company_id, "doc_type": "credit_rating"},
        )
        assert response.status_code == 200
        assert all(d["doc_type"] == "credit_rating" for d in response.json())

    def test_can_exclude_superseded_versions(self, api_client, company_id):
        everything = api_client.get(
            f"{BASE}/documents",
            params={"company_id": company_id, "include_superseded": True},
        ).json()
        current = api_client.get(
            f"{BASE}/documents",
            params={"company_id": company_id, "include_superseded": False},
        ).json()
        assert len(current) <= len(everything)
        assert all(d["superseded_by"] is None for d in current)

    def test_detail_includes_sections_and_pages(self, api_client, uploaded):
        response = api_client.get(f"{BASE}/documents/{uploaded['annual_report']}")
        document = response.json()
        kinds = {s["kind"] for s in document["sections"]}
        assert {"risk_factors", "corporate_governance", "esg"} <= kinds
        assert len(document["pages"]) == document["page_count"]

    def test_page_text_is_retrievable(self, api_client, uploaded):
        response = api_client.get(
            f"{BASE}/documents/{uploaded['annual_report']}/pages/1"
        )
        assert response.status_code == 200
        assert response.json()["text_source"] == "native"
        assert len(response.json()["text"]) > 100

    def test_missing_document_is_404(self, api_client):
        assert api_client.get(f"{BASE}/documents/999999").status_code == 404

    def test_missing_page_is_404(self, api_client, uploaded):
        response = api_client.get(
            f"{BASE}/documents/{uploaded['annual_report']}/pages/9999"
        )
        assert response.status_code == 404


# ===========================================================================
# Chunks, tables, entities, facts
# ===========================================================================
class TestArtefacts:
    def test_chunks_carry_full_provenance(self, api_client, uploaded):
        response = api_client.get(
            f"{BASE}/documents/chunks",
            params={"document_id": uploaded["annual_report"]},
        )
        assert response.status_code == 200
        chunks = response.json()
        assert chunks
        for chunk in chunks:
            assert chunk["page"] >= 1
            assert chunk["fingerprint"]
            assert chunk["section"]

    def test_chunk_pagination(self, api_client, uploaded):
        params = {"document_id": uploaded["annual_report"], "limit": 3}
        first = api_client.get(f"{BASE}/documents/chunks", params=params).json()
        second = api_client.get(
            f"{BASE}/documents/chunks", params={**params, "offset": 3}
        ).json()
        assert len(first) <= 3
        assert not ({c["id"] for c in first} & {c["id"] for c in second})

    def test_tables_preserve_headers_units_and_rows(self, api_client, uploaded):
        response = api_client.get(
            f"{BASE}/documents/tables",
            params={"document_id": uploaded["annual_report"]},
        )
        assert response.status_code == 200
        tables = response.json()
        assert tables
        table = tables[0]
        assert table["unit"] == "inr_cr"
        assert table["header"][0].startswith("Particulars")
        labels = [row[0] for row in table["rows"]]
        assert "Revenue from operations" in labels
        assert table["n_rows"] >= 12

    def test_entities_are_listed_by_kind(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/entities",
            params={"company_id": company_id, "kind": "subsidiary"},
        )
        assert response.status_code == 200
        names = {e["name"] for e in response.json()}
        assert "Bharat Nutrition Private Limited" in names

    def test_entity_confidence_filter(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/entities",
            params={"company_id": company_id, "min_confidence": 0.8},
        )
        assert all(e["confidence"] >= 0.8 for e in response.json())

    def test_entities_require_a_scope(self, api_client):
        assert api_client.get(f"{BASE}/documents/entities").status_code == 400

    def test_facts_reproduce_the_reported_figures(
        self, api_client, company_id, uploaded
    ):
        response = api_client.get(
            f"{BASE}/documents/facts",
            params={"company_id": company_id, "field_key": "revenue"},
        )
        assert response.status_code == 200
        by_period = {f["period"]: f for f in response.json()}
        assert by_period["FY25"]["value"] == pytest.approx(33543.0)
        assert by_period["FY25"]["unit"] == "inr_cr"
        assert by_period["FY24"]["value"] == pytest.approx(28914.0)

    def test_every_fact_cites_a_page(self, api_client, company_id, uploaded):
        facts = api_client.get(
            f"{BASE}/documents/facts", params={"company_id": company_id}
        ).json()
        assert facts
        for fact in facts:
            assert fact["page"] >= 1
            assert fact["evidence"]
            assert 0.0 < fact["confidence"] <= 1.0

    def test_facts_filter_by_category(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/facts",
            params={"company_id": company_id, "category": "GOVERNANCE"},
        )
        assert response.status_code == 200
        assert all(f["category"] == "GOVERNANCE" for f in response.json())

    def test_facts_require_a_scope(self, api_client):
        assert api_client.get(f"{BASE}/documents/facts").status_code == 400


# ===========================================================================
# Search
# ===========================================================================
class TestSearchApi:
    def test_returns_answer_passages_pages_and_confidence(
        self, api_client, company_id, uploaded
    ):
        """The brief's four required elements, in one response."""
        response = api_client.get(
            f"{BASE}/documents/search",
            params={"q": "What is the EBITDA margin guidance?", "company_id": company_id},
        )
        assert response.status_code == 200
        payload = response.json()
        assert "17" in payload["answer"]
        assert payload["hits"]
        assert all(h["page"] >= 1 for h in payload["hits"])
        assert payload["confidence"] > 0.3
        assert payload["took_ms"] > 0

    def test_citations_name_document_page_section_paragraph(
        self, api_client, company_id, uploaded
    ):
        response = api_client.get(
            f"{BASE}/documents/search",
            params={"q": "subsidiaries of the company", "company_id": company_id},
        )
        citations = response.json()["citations"]
        assert citations
        for citation in citations:
            assert set(citation) >= {
                "document_id", "document_title", "page", "section",
                "paragraph", "quote", "reference",
            }
            assert citation["page"] >= 1
            assert citation["quote"]

    def test_citation_audit_verifies_the_answer(
        self, api_client, company_id, uploaded
    ):
        response = api_client.get(
            f"{BASE}/documents/search",
            params={"q": "What is the credit rating?", "company_id": company_id},
        )
        audit = response.json()["citation_audit"]
        assert audit["verified"] is True
        assert audit["unsupported_pages"] == []

    def test_declares_unavailability_instead_of_guessing(
        self, api_client, company_id, uploaded
    ):
        response = api_client.get(
            f"{BASE}/documents/search",
            params={"q": "What is the dividend policy?", "company_id": company_id},
        )
        payload = response.json()
        assert payload["unavailable_reason"] is not None
        assert payload["answer"] == ""

    def test_scopes_to_a_single_document(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/search",
            params={
                "q": "credit rating",
                "company_id": company_id,
                "document_id": [uploaded["rating"]],
            },
        )
        hits = response.json()["hits"]
        assert hits
        assert all(h["document_id"] == uploaded["rating"] for h in hits)

    def test_scopes_to_a_section(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/search",
            params={
                "q": "commodity price volatility",
                "company_id": company_id,
                "section": ["risk_factors"],
            },
        )
        hits = response.json()["hits"]
        assert all(h["section"] == "risk_factors" for h in hits)

    def test_reports_both_score_components(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/search",
            params={"q": "EBITDA margin", "company_id": company_id},
        )
        hit = response.json()["hits"][0]
        assert hit["lexical_score"] >= 0
        assert hit["semantic_score"] >= 0

    def test_empty_query_is_rejected(self, api_client):
        assert api_client.get(
            f"{BASE}/documents/search", params={"q": ""}
        ).status_code == 422


# ===========================================================================
# Knowledge graph
# ===========================================================================
class TestKnowledgeApi:
    def test_returns_nodes_edges_and_stats(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/knowledge", params={"company_id": company_id}
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["stats"]["nodes"] > 5
        assert payload["stats"]["edges"] > 5
        assert payload["company"]["ticker"] == "BHARATCP"

    def test_relationships_cover_the_brief(self, api_client, company_id, uploaded):
        payload = api_client.get(
            f"{BASE}/documents/knowledge", params={"company_id": company_id}
        ).json()
        relations = set(payload["stats"]["relations"])
        assert {"subsidiary_of", "director_of", "competes_with"} <= relations

    def test_every_edge_carries_evidence(self, api_client, company_id, uploaded):
        payload = api_client.get(
            f"{BASE}/documents/knowledge", params={"company_id": company_id}
        ).json()
        for edge in payload["edges"]:
            assert edge["pages"]
            assert edge["confidence"] > 0

    def test_confidence_filter(self, api_client, company_id, uploaded):
        payload = api_client.get(
            f"{BASE}/documents/knowledge",
            params={"company_id": company_id, "min_confidence": 0.75},
        ).json()
        assert all(e["confidence"] >= 0.75 for e in payload["edges"])

    def test_unknown_company_is_404(self, api_client):
        response = api_client.get(
            f"{BASE}/documents/knowledge", params={"company_id": "nope"}
        )
        assert response.status_code == 404


# ===========================================================================
# Coverage, statistics, capabilities
# ===========================================================================
class TestCoverageApi:
    def test_measures_against_the_seventy_three_field_registry(
        self, api_client, company_id, uploaded
    ):
        response = api_client.get(
            f"{BASE}/documents/coverage", params={"company_id": company_id}
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["fields_defined"] == FIELD_COUNT == 73
        assert payload["fields_extracted"] > 30
        assert payload["coverage"] > 0.4
        assert len(payload["categories"]) == 16

    def test_reports_what_was_not_found(self, api_client, company_id, uploaded):
        """A coverage figure without its complement invites a false assumption."""
        payload = api_client.get(
            f"{BASE}/documents/coverage", params={"company_id": company_id}
        ).json()
        for category in payload["categories"]:
            assert category["extracted"] + len(category["missing"]) == \
                category["defined"]

    def test_statistics_counts_the_corpus(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/statistics", params={"company_id": company_id}
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["documents"] >= 3
        assert payload["pages"] > 0
        assert payload["chunks"] > 0
        assert payload["by_type"]["annual_report"] >= 1
        assert payload["embedding"]["dimension"] > 0

    def test_capabilities_describe_the_engine(self, api_client):
        response = api_client.get(f"{BASE}/documents/capabilities")
        assert response.status_code == 200
        payload = response.json()
        assert payload["field_count"] == 73
        assert len(payload["fields"]) == 73
        assert {"pdf", "docx", "txt", "html", "csv", "xlsx"} <= \
            set(payload["file_formats"])
        assert len(payload["pipeline_stages"]) >= 10
        assert "available" in payload["ocr"]

    def test_capabilities_list_every_brief_document_type(self, api_client):
        types = set(api_client.get(f"{BASE}/documents/capabilities").json()[
            "document_types"
        ])
        assert {
            "annual_report", "quarterly_report", "investor_presentation",
            "conference_call", "credit_rating", "drhp", "esg_report",
            "exchange_filing",
        } <= types


# ===========================================================================
# Queue and reindex
# ===========================================================================
class TestOperations:
    def test_jobs_record_per_stage_timings(self, api_client, company_id, uploaded):
        response = api_client.get(
            f"{BASE}/documents/jobs", params={"company_id": company_id}
        )
        assert response.status_code == 200
        jobs = response.json()
        assert jobs
        finished = [j for j in jobs if j["status"] == "ready"]
        assert finished
        assert finished[0]["duration_ms"] > 0
        assert finished[0]["timings"]

    def test_reindex_rebuilds_vectors_without_reparsing(
        self, api_client, company_id, uploaded
    ):
        response = api_client.post(
            f"{BASE}/documents/reindex", params={"company_id": company_id}
        )
        assert response.status_code == 200
        payload = response.json()
        assert payload["reindexed_chunks"] > 0
        assert payload["embedding"]["dimension"] > 0

    def test_search_still_works_after_reindex(self, api_client, company_id, uploaded):
        api_client.post(f"{BASE}/documents/reindex", params={"company_id": company_id})
        response = api_client.get(
            f"{BASE}/documents/search",
            params={"q": "What is the credit rating?", "company_id": company_id},
        )
        assert "AA+" in response.json()["answer"]

    def test_delete_removes_the_document(self, api_client, company_id):
        created = api_client.post(
            f"{BASE}/documents/upload",
            data={"company_id": company_id},
            files={"file": ("disposable.txt", b"Revenue was Rs 100 crore.", "text/plain")},
        ).json()["document"]["id"]
        assert api_client.delete(f"{BASE}/documents/{created}").status_code == 204
        assert api_client.get(f"{BASE}/documents/{created}").status_code == 404

    def test_delete_of_a_missing_document_is_404(self, api_client):
        assert api_client.delete(f"{BASE}/documents/999999").status_code == 404


# ===========================================================================
# Module 6 integration — the AI layer now sees documents
# ===========================================================================
class TestAiIntegration:
    def test_ai_context_now_includes_document_evidence(
        self, api_client, uploaded
    ):
        """Module 6 shipped with EvidenceKind.DOCUMENT permanently empty.

        Module 7's whole purpose on that axis is to fill it. If this returns
        nothing, the two modules are not actually connected.
        """
        response = api_client.get(f"{BASE}/company/BHARATCP/ai/context")
        assert response.status_code == 200
        payload = response.json()
        citations = payload.get("citations", [])
        document_citations = [c for c in citations if c["kind"] == "document"]
        assert document_citations, "no document evidence reached the AI context"

    def test_document_citations_name_their_source_page(self, api_client, uploaded):
        payload = api_client.get(f"{BASE}/company/BHARATCP/ai/context").json()
        for citation in payload["citations"]:
            if citation["kind"] == "document":
                assert "p." in citation["source"]

    def test_analysis_still_runs_with_document_evidence_present(
        self, api_client, uploaded
    ):
        response = api_client.post(
            f"{BASE}/company/BHARATCP/ai/analyse",
            json={"capability": "business_summary"},
        )
        assert response.status_code == 200
        assert response.json()["content"]


# ===========================================================================
# Module 6 tool wiring — document_search is now a real retrieval
# ===========================================================================
class TestDocumentSearchTool:
    """Module 6 shipped `document_search` as a placeholder.

    It returned the same pre-built context regardless of the question, which
    was honest only because nothing populated document evidence. Now that
    documents exist, the tool must genuinely retrieve — a tool whose answer
    does not depend on its argument is not a search.
    """

    def _runner(self, api_client, uploaded):
        from app.services.ai.service import AIService
        from app.services.ai.tools import ToolRunner
        from app.services.analysis_service import AnalysisService
        from tests.conftest import TestingSession

        db = TestingSession()
        analysis = AnalysisService.for_ticker(db, "BHARATCP")
        analyst = AIService(db).analyst_for(analysis)
        builder = getattr(analyst, "builder", None) or analyst.context_builder
        return ToolRunner(builder), db

    def test_tool_returns_document_citations(self, api_client, uploaded):
        runner, db = self._runner(api_client, uploaded)
        try:
            citations = runner.run("document_search", query="EBITDA margin guidance")
            assert citations
            assert all(c.kind.value == "document" for c in citations)
            assert all("page" in c.source for c in citations)
        finally:
            db.close()

    def test_results_depend_on_the_query(self, api_client, uploaded):
        runner, db = self._runner(api_client, uploaded)
        try:
            rating = runner.run("document_search", query="credit rating covenants")
            risks = runner.run("document_search", query="agricultural commodity risk")
            assert [c.value for c in rating] != [c.value for c in risks]
            assert any("CRISIL" in str(c.value) or "CRISIL" in c.source for c in rating)
        finally:
            db.close()

    def test_citation_value_is_the_quoted_text(self, api_client, uploaded):
        """A passage citation must carry the words, not a relevance score.

        Otherwise the model has nothing to quote and everything to invent.
        """
        runner, db = self._runner(api_client, uploaded)
        try:
            for citation in runner.run("document_search", query="subsidiaries"):
                assert isinstance(citation.value, str)
                assert len(citation.value) > 20
        finally:
            db.close()

    def test_other_tools_are_unaffected(self, api_client, uploaded):
        runner, db = self._runner(api_client, uploaded)
        try:
            ratios = runner.run("ratio_lookup")
            assert ratios and all(c.kind.value == "ratio" for c in ratios)
        finally:
            db.close()
