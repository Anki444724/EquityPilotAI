"""Parsers for DOCX, TXT, Markdown, HTML, CSV and XLSX.

Each reduces to the same :class:`ParsedDocument`. The only real design question
is what "page" means for a format that has none.

* **DOCX** — Word stores no page breaks that survive without a layout engine,
  so pages are synthesised at a fixed paragraph count. Arbitrary, but stable
  and reproducible, which is all a citation ordinal needs to be.
* **HTML / TXT / MD** — paginated the same way, on block boundaries.
* **CSV** — one page, one table. The whole file *is* the table.
* **XLSX** — one page per worksheet, which is the only mapping a user would
  recognise when told "page 3".

None of these ever need OCR, so the OCR engine is not even imported here.
"""
from __future__ import annotations

import csv
import io
import re
from typing import ClassVar

from app.domain.documents.types import (
    FileFormat, ParseFailure, ParsedDocument, ParsedPage, TextBlock, TextSource,
    Unit, normalise_whitespace,
)
from app.services.documents.extractors.base import DocumentParser, register
from app.services.documents.extractors.tables import build_table, detect_unit

#: Blocks per synthesised page for formats with no intrinsic pagination.
BLOCKS_PER_PAGE = 40


def _paginate(blocks: list[TextBlock]) -> list[ParsedPage]:
    """Group blocks into synthetic pages of :data:`BLOCKS_PER_PAGE`."""
    pages: list[ParsedPage] = []
    if not blocks:
        return [ParsedPage(number=1, text="", source=TextSource.NATIVE)]
    for start in range(0, len(blocks), BLOCKS_PER_PAGE):
        number = start // BLOCKS_PER_PAGE + 1
        window = blocks[start : start + BLOCKS_PER_PAGE]
        for block in window:
            block.page = number
        pages.append(
            ParsedPage(
                number=number,
                text="\n".join(b.text for b in window if not b.is_empty),
                blocks=window,
                source=TextSource.NATIVE,
            )
        )
    return pages


def _decode(payload: bytes) -> str:
    """Decode text, tolerating the encodings Indian filings actually arrive in."""
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
@register
class DocxParser(DocumentParser):
    """Word documents, including their tables."""

    formats: ClassVar[tuple[FileFormat, ...]] = (FileFormat.DOCX,)

    def parse(self, payload: bytes, *, filename: str = "") -> ParsedDocument:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ParseFailure("python-docx is not installed") from exc
        try:
            document = docx.Document(io.BytesIO(payload))
        except Exception as exc:
            raise ParseFailure(f"cannot open DOCX: {exc}") from exc

        blocks: list[TextBlock] = []
        for index, paragraph in enumerate(document.paragraphs):
            text = normalise_whitespace(paragraph.text)
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style else "") or ""
            heading = style.lower().startswith("heading") or style.lower() == "title"
            blocks.append(
                TextBlock(
                    text=text,
                    page=1,
                    # Word carries semantic styles, so heading detection here is
                    # a fact rather than the font-size inference a PDF forces.
                    font_size=18.0 if heading else 11.0,
                    bold=heading or self._mostly_bold(paragraph),
                    block_index=index,
                )
            )

        parsed = ParsedDocument(file_format=FileFormat.DOCX)
        parsed.pages = _paginate(blocks)
        self._attach_tables(document, parsed)

        properties = getattr(document, "core_properties", None)
        if properties is not None:
            parsed.title = (properties.title or "").strip() or None
            parsed.author = (properties.author or "").strip() or None
        return parsed

    @staticmethod
    def _mostly_bold(paragraph) -> bool:
        runs = [r for r in paragraph.runs if r.text.strip()]
        if not runs:
            return False
        bold = sum(len(r.text) for r in runs if r.bold)
        total = sum(len(r.text) for r in runs)
        return total > 0 and bold / total > 0.6

    @staticmethod
    def _attach_tables(document, parsed: ParsedDocument) -> None:
        if not parsed.pages:
            return
        for index, table in enumerate(document.tables):
            raw = [[cell.text for cell in row.cells] for row in table.rows]
            # Word tables are not anchored to a synthetic page, so they attach
            # to the first — stated plainly rather than fabricating a location.
            built = build_table(raw, page=1, index=index)
            if built is not None:
                parsed.pages[0].tables.append(built)


# ---------------------------------------------------------------------------
@register
class PlainTextParser(DocumentParser):
    """TXT and Markdown."""

    formats: ClassVar[tuple[FileFormat, ...]] = (FileFormat.TXT, FileFormat.MARKDOWN)

    _ATX = re.compile(r"^(#{1,6})\s+(.*)$")
    _SETEXT = re.compile(r"^(=+|-{3,})\s*$")

    def parse(self, payload: bytes, *, filename: str = "") -> ParsedDocument:
        text = _decode(payload)
        blocks: list[TextBlock] = []
        title: str | None = None

        for index, raw in enumerate(re.split(r"\n\s*\n", text)):
            content = raw.strip("\n")
            if not content.strip():
                continue
            first = content.splitlines()[0]
            heading = self._ATX.match(first.strip())
            lines = content.splitlines()
            setext = len(lines) > 1 and bool(self._SETEXT.match(lines[1].strip()))
            is_heading = bool(heading) or setext
            if is_heading and title is None:
                title = normalise_whitespace(
                    heading.group(2) if heading else lines[0]
                )
            blocks.append(
                TextBlock(
                    text=normalise_whitespace(content.replace("\n", " "))
                    if is_heading
                    else content.strip(),
                    page=1,
                    font_size=18.0 if is_heading else 11.0,
                    bold=is_heading,
                    block_index=index,
                )
            )

        parsed = ParsedDocument(file_format=FileFormat.TXT)
        parsed.pages = _paginate(blocks)
        parsed.title = title
        return parsed


# ---------------------------------------------------------------------------
@register
class HtmlParser(DocumentParser):
    """HTML, including exchange-filing pages saved from a browser."""

    formats: ClassVar[tuple[FileFormat, ...]] = (FileFormat.HTML,)

    _HEADINGS = {"h1": 26.0, "h2": 22.0, "h3": 18.0, "h4": 15.0, "h5": 13.0, "h6": 12.0}

    def parse(self, payload: bytes, *, filename: str = "") -> ParsedDocument:
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ParseFailure("beautifulsoup4 is not installed") from exc

        soup = BeautifulSoup(_decode(payload), "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        blocks: list[TextBlock] = []
        tables: list = []
        index = 0
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "div", "table"]
        ):
            if element.name == "table":
                raw = [
                    [cell.get_text(" ", strip=True) for cell in row.find_all(["td", "th"])]
                    for row in element.find_all("tr")
                ]
                built = build_table(raw, page=1, index=len(tables))
                if built is not None:
                    tables.append(built)
                continue
            if element.find(["p", "div", "table", "li"]):
                continue  # container; its children carry the text
            text = normalise_whitespace(element.get_text(" ", strip=True))
            if not text:
                continue
            size = self._HEADINGS.get(element.name, 11.0)
            blocks.append(
                TextBlock(
                    text=text,
                    page=1,
                    font_size=size,
                    bold=element.name in self._HEADINGS,
                    block_index=index,
                )
            )
            index += 1

        parsed = ParsedDocument(file_format=FileFormat.HTML)
        parsed.pages = _paginate(blocks)
        if tables and parsed.pages:
            parsed.pages[0].tables.extend(tables)
        if soup.title and soup.title.string:
            parsed.title = normalise_whitespace(soup.title.string)
        return parsed


# ---------------------------------------------------------------------------
@register
class CsvParser(DocumentParser):
    """CSV — the file is one table, so it becomes one page holding one table."""

    formats: ClassVar[tuple[FileFormat, ...]] = (FileFormat.CSV,)

    def parse(self, payload: bytes, *, filename: str = "") -> ParsedDocument:
        text = _decode(payload)
        try:
            dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel  # a single-column file sniffs as nothing
        rows = [row for row in csv.reader(io.StringIO(text), dialect)]
        if not rows:
            raise ParseFailure("CSV contains no rows")

        table = build_table(rows, page=1, index=0, caption=filename or None)
        rendered = "\n".join(" | ".join(cell for cell in row) for row in rows)
        page = ParsedPage(
            number=1,
            text=rendered,
            source=TextSource.STRUCTURED,
            tables=[table] if table else [],
        )
        parsed = ParsedDocument(file_format=FileFormat.CSV, pages=[page])
        parsed.title = filename or None
        return parsed


# ---------------------------------------------------------------------------
@register
class ExcelParser(DocumentParser):
    """XLSX — one page per worksheet, each worksheet one table."""

    formats: ClassVar[tuple[FileFormat, ...]] = (FileFormat.XLSX,)

    def parse(self, payload: bytes, *, filename: str = "") -> ParsedDocument:
        try:
            import openpyxl
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ParseFailure("openpyxl is not installed") from exc
        try:
            workbook = openpyxl.load_workbook(
                io.BytesIO(payload), read_only=True, data_only=True
            )
        except Exception as exc:
            raise ParseFailure(f"cannot open workbook: {exc}") from exc

        parsed = ParsedDocument(file_format=FileFormat.XLSX)
        try:
            for number, name in enumerate(workbook.sheetnames, start=1):
                sheet = workbook[name]
                rows = [
                    ["" if cell is None else str(cell) for cell in row]
                    for row in sheet.iter_rows(values_only=True)
                ]
                rows = [row for row in rows if any(cell.strip() for cell in row)]
                table = build_table(rows, page=number, index=0, caption=name)
                # Excel declares merges natively; prefer that over inference.
                if table is not None:
                    declared = self._declared_merges(sheet)
                    if declared:
                        table.merged.update(declared)
                    if table.unit is Unit.UNKNOWN:
                        table.unit = detect_unit(name)
                parsed.pages.append(
                    ParsedPage(
                        number=number,
                        text="\n".join(" | ".join(row) for row in rows),
                        source=TextSource.STRUCTURED,
                        tables=[table] if table else [],
                    )
                )
        finally:
            workbook.close()

        if not parsed.pages:
            raise ParseFailure("workbook contains no readable sheets")
        parsed.title = filename or None
        return parsed

    @staticmethod
    def _declared_merges(sheet) -> dict[tuple[int, int], tuple[int, int]]:
        """Read openpyxl's merged ranges, zero-indexed to match our grid."""
        merges: dict[tuple[int, int], tuple[int, int]] = {}
        for cell_range in getattr(sheet, "merged_cells", None) or []:
            try:
                bounds = cell_range.bounds  # (min_col, min_row, max_col, max_row)
            except AttributeError:  # pragma: no cover - read-only mode variance
                continue
            min_col, min_row, max_col, max_row = bounds
            merges[(min_row - 1, min_col - 1)] = (
                max_row - min_row + 1,
                max_col - min_col + 1,
            )
        return merges
