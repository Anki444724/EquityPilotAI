"""Word renderer.

python-docx, producing a real `.docx` — the modern OOXML format, which previews
in-app. A legacy `.doc` would only download.

Word's own heading styles are used rather than hand-formatted bold text,
because that is what populates the navigation pane and lets Word build a
field-driven table of contents. A report whose headings are merely large bold
paragraphs looks identical and is unnavigable.
"""
from __future__ import annotations

import io
import time
from typing import Any

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.domain.reports.blocks import (
    Block, BlockKind, Callout, CalloutTone, Chart, ReportDocument, SectionKey,
    Table, Theme,
)
from app.domain.reports.citations import strip_markers
from app.services.reports.renderers.base import (
    OutputFormat, RenderResult, ReportRenderer, cover_pairs, register,
)

NAVY = RGBColor(0x1E, 0x3A, 0x8A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
POSITIVE = RGBColor(0x05, 0x96, 0x69)
NEGATIVE = RGBColor(0xDC, 0x26, 0x26)
WARNING = RGBColor(0xB4, 0x53, 0x09)

TONE_COLOURS = {
    CalloutTone.POSITIVE: POSITIVE,
    CalloutTone.NEGATIVE: NEGATIVE,
    CalloutTone.WARNING: WARNING,
    CalloutTone.NEUTRAL: NAVY,
}


@register
class DocxRenderer(ReportRenderer):
    """Renders a report to a Word document."""

    fmt = OutputFormat.DOCX

    def render(self, document: ReportDocument) -> RenderResult:
        started = time.perf_counter()
        docx = DocxDocument()
        self._page_setup(docx)
        self._styles(docx)
        engine = self.chart_engine(document.theme)
        evidence = document.evidence()

        self._cover(docx, document)
        docx.add_page_break()

        for section in document.ordered():
            if section.key is SectionKey.COVER:
                continue
            if section.key is SectionKey.TOC:
                self._toc(docx)
                docx.add_page_break()
                continue
            docx.add_heading(section.title, level=1)
            for block in section.blocks:
                self._block(docx, block, engine, evidence)

        self._footer(docx, document)

        buffer = io.BytesIO()
        docx.save(buffer)
        return RenderResult(
            payload=buffer.getvalue(), fmt=self.fmt,
            filename=self.filename(document, self.fmt),
            took_ms=round((time.perf_counter() - started) * 1000, 2),
        )

    # ------------------------------------------------------------------
    @staticmethod
    def _page_setup(docx) -> None:
        for section in docx.sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)

    @staticmethod
    def _styles(docx) -> None:
        normal = docx.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(9.5)
        for level, size in ((1, 14), (2, 11.5), (3, 10)):
            style = docx.styles[f"Heading {level}"]
            style.font.size = Pt(size)
            style.font.color.rgb = NAVY if level == 1 else RGBColor(0x0F, 0x17, 0x2A)
            style.font.name = "Calibri"

    def _cover(self, docx, document: ReportDocument) -> None:
        cover = document.cover
        institution = docx.add_paragraph()
        run = institution.add_run(cover.institution.upper())
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        run.bold = True

        title = docx.add_paragraph()
        run = title.add_run(cover.company_name)
        run.font.size = Pt(26)
        run.bold = True
        run.font.color.rgb = NAVY

        subtitle = docx.add_paragraph()
        run = subtitle.add_run(f"{cover.ticker} · {cover.title}")
        run.font.size = Pt(12)
        run.font.color.rgb = MUTED

        docx.add_paragraph()
        table = docx.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for key, value in cover_pairs(document):
            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = value
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        if cover.data_warning:
            docx.add_paragraph()
            self._callout(
                docx,
                Callout("Data quality", cover.data_warning, CalloutTone.WARNING),
            )

        docx.add_paragraph()
        disclaimer = docx.add_paragraph()
        run = disclaimer.add_run(document.disclaimer)
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED

    @staticmethod
    def _toc(docx) -> None:
        """Insert a Word TOC field.

        A field rather than a rendered list: Word computes the page numbers
        itself and refreshes them if the document is edited. A static list
        would be wrong the moment anyone touched the file.
        """
        docx.add_heading("Contents", level=1)
        paragraph = docx.add_paragraph()
        run = paragraph.add_run()

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = r'TOC \o "1-3" \h \z \u'
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = (
            "Right-click and choose 'Update Field' to build the contents."
        )
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        for element in (begin, instruction, separate, placeholder, end):
            run._r.append(element)

    @staticmethod
    def _footer(docx, document: ReportDocument) -> None:
        """Page numbers via a PAGE field, so Word keeps them accurate."""
        footer = docx.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover = document.cover
        run = paragraph.add_run(
            f"{cover.company_name} ({cover.ticker}) · Not investment advice · Page "
        )
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED

        field_run = paragraph.add_run()
        field_run.font.size = Pt(7.5)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for element in (begin, instruction, end):
            field_run._r.append(element)

    # ------------------------------------------------------------------
    def _block(self, docx, block: Block, engine, evidence) -> None:
        kind = block.kind

        if kind is BlockKind.HEADING:
            docx.add_heading(block.text, level=min(3, max(2, block.level)))

        elif kind is BlockKind.PARAGRAPH:
            self._prose(docx, block.text, evidence)

        elif kind is BlockKind.BULLETS:
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                docx.add_paragraph(strip_markers(item), style=style)

        elif kind is BlockKind.KEY_VALUE:
            table = docx.add_table(rows=0, cols=2)
            table.style = "Light Grid Accent 1"
            for key, value in block.pairs:
                cells = table.add_row().cells
                cells[0].text = key
                cells[1].text = value

        elif kind is BlockKind.TABLE:
            self._table(docx, block)

        elif kind is BlockKind.METRIC_GRID:
            columns = max(1, min(block.columns, 4))
            table = docx.add_table(rows=0, cols=columns)
            table.style = "Light Grid Accent 1"
            row = None
            for index, (label, value, hint) in enumerate(block.metrics):
                if index % columns == 0:
                    row = table.add_row()
                cell = row.cells[index % columns]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                label_run = paragraph.add_run(f"{label.upper()}\n")
                label_run.font.size = Pt(7)
                label_run.font.color.rgb = MUTED
                value_run = paragraph.add_run(f"{value}\n")
                value_run.font.size = Pt(13)
                value_run.bold = True
                hint_run = paragraph.add_run(hint)
                hint_run.font.size = Pt(7)
                hint_run.font.color.rgb = MUTED

        elif kind is BlockKind.CHART:
            png = engine.render(block)
            if png is not None:
                docx.add_picture(io.BytesIO(png), width=Inches(6.6))

        elif kind is BlockKind.CALLOUT:
            self._callout(docx, block)

        elif kind is BlockKind.QUOTE:
            paragraph = docx.add_paragraph(block.text, style="Intense Quote")
            if block.attribution:
                run = paragraph.add_run(f"\n— {block.attribution}")
                run.italic = True
                run.font.size = Pt(8)

        elif kind is BlockKind.DIVIDER:
            docx.add_paragraph("─" * 60).runs[0].font.color.rgb = MUTED

        elif kind is BlockKind.PAGE_BREAK:
            docx.add_page_break()

        elif kind is BlockKind.INSUFFICIENT:
            self._callout(
                docx,
                Callout("Insufficient evidence", block.reason or "",
                        CalloutTone.NEUTRAL),
            )

        elif kind is BlockKind.CITATION_LIST:
            if block.entries:
                self._table(docx, Table(
                    ["Reference", "Evidence"],
                    [[e.key, e.render()] for e in block.entries],
                ))

        if block.note:
            run = docx.add_paragraph().add_run(block.note)
            run.font.size = Pt(7.5)
            run.font.color.rgb = MUTED

    def _prose(self, docx, text: str, evidence) -> None:
        """Render prose with citation markers as small grey superscripts."""
        import re

        lookup = {e.key: e for e in evidence}
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cursor = 0
        for match in re.finditer(r"\[([a-z0-9_]+)\]", text):
            paragraph.add_run(text[cursor:match.start()])
            found = lookup.get(match.group(1))
            if found is None:
                # An unresolvable marker stays visible. Hiding it would make a
                # broken reference read as though it were never cited.
                paragraph.add_run(match.group(0))
            else:
                run = paragraph.add_run(f" [{found.label}]")
                run.font.size = Pt(6.5)
                run.font.color.rgb = NAVY
            cursor = match.end()
        paragraph.add_run(text[cursor:])

    @staticmethod
    def _table(docx, block: Table) -> None:
        if not block.header and not block.rows:
            return
        if block.caption:
            docx.add_heading(block.caption, level=3)
        columns = block.n_cols or 1
        table = docx.add_table(rows=1 if block.header else 0, cols=columns)
        table.style = "Light Grid Accent 1"

        if block.header:
            for index, heading in enumerate(block.header[:columns]):
                cell = table.rows[0].cells[index]
                cell.text = heading
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8)

        for row_index, row in enumerate(block.rows):
            cells = table.add_row().cells
            for index, value in enumerate(row[:columns]):
                cells[index].text = value
                align = block.align[index] if index < len(block.align) else "l"
                for paragraph in cells[index].paragraphs:
                    if align == "r":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        if row_index in block.emphasis_rows:
                            run.bold = True

    @staticmethod
    def _callout(docx, block: Callout) -> None:
        paragraph = docx.add_paragraph()
        title = paragraph.add_run(f"{block.title}\n")
        title.bold = True
        title.font.color.rgb = TONE_COLOURS[block.tone]
        body = paragraph.add_run(block.text)
        body.font.size = Pt(9)

        # A shaded background, since python-docx has no border API for
        # paragraphs. Applied as raw OOXML.
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), "F5F7FA")
        paragraph._p.get_or_add_pPr().append(shading)
